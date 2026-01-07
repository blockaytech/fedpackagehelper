#!/usr/bin/env python3
"""
FedRAMP Package Manager

A unified tool for managing FedRAMP package documentation updates.

Workflow:
    1. Place original documents in 'originals/' directory
    2. Run 'analyze' to discover terms
    3. Run 'preview' to review changes
    4. Run 'apply' to create drafts with changes (originals preserved)
    5. Run 'verify' to validate drafts and check completeness

Commands:
    status      - Show current configuration status
    analyze     - Scan documents and discover terms
    preview     - Preview replacements across all documents
    apply       - Apply replacements (originals -> drafts workflow)
    verify      - Re-analyze drafts to verify completeness
    export      - Export analysis to Excel/CSV

Usage:
    python package_manager.py status
    python package_manager.py analyze
    python package_manager.py preview
    python package_manager.py apply
    python package_manager.py verify
    python package_manager.py export --format excel
"""

import re
import json
import argparse
import shutil
from datetime import datetime
from pathlib import Path
from collections import defaultdict
from copy import deepcopy

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# Optional imports for export
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


# =============================================================================
# CONFIGURATION
# =============================================================================

class Config:
    """Central configuration for the package manager."""

    def __init__(self, base_dir: str = "."):
        self.base_dir = Path(base_dir)
        self.terms_dict_path = self.base_dir / "terms_dictionary.json"
        self.replacements_path = self.base_dir / "replacements.json"

        # Directory structure for workflow
        self.originals_dir = self.base_dir / "originals"   # Source documents (never modified)
        self.drafts_dir = self.base_dir / "drafts"         # Working copies with changes
        self.output_dir = self.base_dir / "output"         # Reports and exports
        self.backup_dir = self.base_dir / "backups"        # Historical backups

        # Processing options
        self.whole_word_matching = True
        self.case_sensitive = True

        # File patterns to process
        self.include_patterns = ["*.docx"]
        self.exclude_patterns = ["~$*", "*_BACKUP_*", "*_UPDATED_*"]

    def get_originals(self) -> list:
        """Get all original documents to process."""
        if not self.originals_dir.exists():
            return []

        docs = []
        for pattern in self.include_patterns:
            docs.extend(self.originals_dir.glob(pattern))

        return self._filter_docs(docs)

    def get_drafts(self) -> list:
        """Get all draft documents."""
        if not self.drafts_dir.exists():
            return []

        docs = []
        for pattern in self.include_patterns:
            docs.extend(self.drafts_dir.glob(pattern))

        return self._filter_docs(docs)

    def get_documents(self) -> list:
        """Get all documents to process (checks originals first, then base dir)."""
        # First check originals directory
        if self.originals_dir.exists():
            docs = self.get_originals()
            if docs:
                return docs

        # Fall back to base directory for backward compatibility
        docs = []
        for pattern in self.include_patterns:
            docs.extend(self.base_dir.glob(pattern))

        return self._filter_docs(docs)

    def _filter_docs(self, docs: list) -> list:
        """Filter documents based on exclusion patterns."""
        filtered = []
        for doc in docs:
            excluded = False
            for exc_pattern in self.exclude_patterns:
                if doc.match(exc_pattern):
                    excluded = True
                    break
            if not excluded:
                filtered.append(doc)

        return sorted(filtered)

    def ensure_dirs(self):
        """Create output directories if needed."""
        self.originals_dir.mkdir(exist_ok=True)
        self.drafts_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)
        self.backup_dir.mkdir(exist_ok=True)


# =============================================================================
# TERMS DICTIONARY
# =============================================================================

class TermsDictionary:
    """Manages the master terms dictionary."""

    def __init__(self, dict_path: Path):
        self.dict_path = dict_path
        self.data = self._load()

    def _load(self) -> dict:
        if not self.dict_path.exists():
            return self._default_dict()
        with open(self.dict_path, 'r') as f:
            return json.load(f)

    def _default_dict(self) -> dict:
        return {
            "known_technologies": {"terms": {}},
            "known_teams": {"terms": {}},
            "known_positions": {"terms": {}},
            "discovery_patterns": {},
            "exclusions": {"terms": []}
        }

    def get_all_known_terms(self) -> dict:
        terms = {}
        for term, info in self.data.get("known_technologies", {}).get("terms", {}).items():
            terms[term] = {"type": "technology", "category": info.get("category", "unknown")}
        for term, info in self.data.get("known_teams", {}).get("terms", {}).items():
            terms[term] = {"type": "team"}
        for term, info in self.data.get("known_positions", {}).get("terms", {}).items():
            terms[term] = {"type": "position", "acronym": info.get("acronym")}
        return terms

    def get_exclusions(self) -> list:
        return self.data.get("exclusions", {}).get("terms", [])


# =============================================================================
# REPLACEMENTS
# =============================================================================

class ReplacementsConfig:
    """Manages replacements configuration."""

    def __init__(self, replacements_path: Path):
        self.path = replacements_path
        self.replacements = self._load()

    def _load(self) -> dict:
        if not self.path.exists():
            return {}

        with open(self.path, 'r') as f:
            data = json.load(f)

        replacements = data.get("replacements", {})

        # Filter and process
        filtered = {}
        for k, v in replacements.items():
            if k.startswith("=====") or v == "DELETE_THIS_LINE" or k.startswith("_"):
                continue
            if v in ("DELETE", "REMOVE"):
                filtered[k] = ""
            else:
                filtered[k] = v

        return filtered

    def get_ordered(self) -> list:
        """Get replacements ordered by length (longest first)."""
        return sorted(self.replacements.items(), key=lambda x: len(x[0]), reverse=True)


# =============================================================================
# TEXT QUALITY CHECKS
# =============================================================================

def find_repeated_words(text: str, min_word_length: int = 2) -> list:
    """
    Find adjacent repeated words in text (e.g., 'Knox Knox').

    Returns list of dicts with:
        - word: the repeated word
        - count: how many times it repeats consecutively
        - context: surrounding text for reference
        - position: character position in text
    """
    if not text:
        return []

    # Pattern to find repeated words (case-insensitive)
    # Matches: word followed by whitespace and the same word
    pattern = re.compile(
        r'\b(\w{' + str(min_word_length) + r',})\s+\1\b',
        re.IGNORECASE
    )

    results = []
    for match in pattern.finditer(text):
        word = match.group(1)
        position = match.start()

        # Get context (50 chars before and after)
        context_start = max(0, position - 30)
        context_end = min(len(text), match.end() + 30)
        context = text[context_start:context_end]

        # Check for more than 2 repetitions (e.g., "Knox Knox Knox")
        full_match = match.group(0)
        extended_pattern = re.compile(
            rf'\b({re.escape(word)}(?:\s+{re.escape(word)})+)\b',
            re.IGNORECASE
        )
        extended_match = extended_pattern.search(text[position:position + 200])
        if extended_match:
            full_match = extended_match.group(1)
            repetition_count = len(re.findall(rf'\b{re.escape(word)}\b', full_match, re.IGNORECASE))
        else:
            repetition_count = 2

        results.append({
            "word": word,
            "count": repetition_count,
            "context": context.strip(),
            "position": position,
            "full_match": full_match
        })

    return results


def fix_repeated_words(text: str, min_word_length: int = 2) -> tuple:
    """
    Fix adjacent repeated words in text (e.g., 'Knox Knox' -> 'Knox').

    Returns tuple of (fixed_text, list of fixes made).
    Each fix is a dict with: word, original, position
    """
    if not text:
        return text, []

    fixes = []

    # Pattern to find repeated words (case-insensitive)
    # This handles 2+ repetitions: "Knox Knox" or "Knox Knox Knox"
    def replace_repeated(match):
        word = match.group(1)
        full_match = match.group(0)
        fixes.append({
            "word": word,
            "original": full_match,
            "position": match.start()
        })
        return word  # Keep just one instance

    # Pattern matches: word followed by one or more repetitions of whitespace + same word
    pattern = re.compile(
        r'\b(\w{' + str(min_word_length) + r',})(\s+\1)+\b',
        re.IGNORECASE
    )

    fixed_text = pattern.sub(replace_repeated, text)

    return fixed_text, fixes


def check_document_for_repeated_words(doc_path: Path, min_word_length: int = 2) -> dict:
    """
    Check a document for repeated adjacent words.

    Returns dict with:
        - document: filename
        - issues: list of repeated word findings with location info
        - total_issues: count of issues found
    """
    from docx import Document

    doc = Document(str(doc_path))
    issues = []

    def check_paragraph(para_text: str, location: str):
        repeated = find_repeated_words(para_text, min_word_length)
        for r in repeated:
            issues.append({
                "location": location,
                "word": r["word"],
                "count": r["count"],
                "context": r["context"],
                "full_match": r["full_match"]
            })

    # Check body paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        if para.text.strip():
            check_paragraph(para.text, f"Body, Paragraph {para_idx + 1}")

    # Check tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    if para.text.strip():
                        location = f"Table {table_idx + 1}, Row {row_idx + 1}, Col {col_idx + 1}"
                        check_paragraph(para.text, location)

    # Check headers/footers
    for section_idx, section in enumerate(doc.sections):
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    if para.text.strip():
                        check_paragraph(para.text, f"Header (Section {section_idx + 1})")
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    if para.text.strip():
                        check_paragraph(para.text, f"Footer (Section {section_idx + 1})")

    return {
        "document": doc_path.name,
        "issues": issues,
        "total_issues": len(issues)
    }


# =============================================================================
# DOCUMENT PROCESSOR
# =============================================================================

class DocumentProcessor:
    """Processes a single document for replacements."""

    def __init__(self, doc_path: Path, config: Config):
        self.doc_path = doc_path
        self.config = config
        self.changes = []
        self.summary = defaultdict(int)

    def extract_text(self, doc: Document) -> str:
        """Extract all text from document."""
        all_text = []
        for para in doc.paragraphs:
            all_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text.append(para.text)
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        all_text.append(para.text)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        all_text.append(para.text)
        return "\n".join(all_text)

    def create_pattern(self, text: str) -> re.Pattern:
        """Create regex pattern for matching."""
        escaped = re.escape(text)
        if self.config.whole_word_matching:
            pattern = rf'\b{escaped}\b'
        else:
            pattern = escaped
        flags = 0 if self.config.case_sensitive else re.IGNORECASE
        return re.compile(pattern, flags)

    def process_paragraph(self, paragraph: Paragraph, replacements: list,
                          location: str, preview_only: bool) -> bool:
        """Process a single paragraph."""
        full_text = paragraph.text
        if not full_text.strip():
            return False

        changes_made = False

        # Calculate new text
        new_full_text = full_text
        for old_text, new_text in replacements:
            pattern = self.create_pattern(old_text)
            new_full_text = pattern.sub(new_text, new_full_text)

        # Log changes
        for old_text, new_text in replacements:
            pattern = self.create_pattern(old_text)
            if pattern.search(full_text):
                for match in pattern.finditer(full_text):
                    context_start = max(0, match.start() - 50)
                    context_end = min(len(full_text), match.end() + 50)
                    context_before = full_text[context_start:context_end]

                    self.changes.append({
                        "document": self.doc_path.name,
                        "location": location,
                        "old_text": old_text,
                        "new_text": new_text if new_text else "[DELETED]",
                        "context_before": context_before,
                        "is_deletion": new_text == ""
                    })

                    key = f"{old_text} -> {new_text if new_text else '[DELETED]'}"
                    self.summary[key] += 1
                    changes_made = True

        # Apply changes
        if changes_made and not preview_only:
            if paragraph.runs:
                for i, run in enumerate(paragraph.runs):
                    run_text = run.text
                    for old_text, new_text in replacements:
                        pattern = self.create_pattern(old_text)
                        run_text = pattern.sub(new_text, run_text)
                    # Fix any repeated words created by replacements (e.g., "Knox Knox" -> "Knox")
                    run_text, _ = fix_repeated_words(run_text)
                    run.text = run_text
            else:
                # Fix repeated words in full text too
                new_full_text, _ = fix_repeated_words(new_full_text)
                paragraph.text = new_full_text

        return changes_made

    def process(self, replacements: list, preview_only: bool = True) -> dict:
        """Process the document with the given replacements."""
        doc = Document(str(self.doc_path))
        changes_made = False

        # Body paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            location = f"Body, Paragraph {para_idx + 1}"
            if self.process_paragraph(paragraph, replacements, location, preview_only):
                changes_made = True

        # Tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        location = f"Table {table_idx + 1}, Row {row_idx + 1}, Col {col_idx + 1}"
                        if self.process_paragraph(paragraph, replacements, location, preview_only):
                            changes_made = True

        # Headers and footers
        for section_idx, section in enumerate(doc.sections):
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para_idx, paragraph in enumerate(header.paragraphs):
                        location = f"Header (Section {section_idx + 1})"
                        if self.process_paragraph(paragraph, replacements, location, preview_only):
                            changes_made = True
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para_idx, paragraph in enumerate(footer.paragraphs):
                        location = f"Footer (Section {section_idx + 1})"
                        if self.process_paragraph(paragraph, replacements, location, preview_only):
                            changes_made = True

        return {
            "document": self.doc_path.name,
            "changes_made": changes_made,
            "change_count": len(self.changes),
            "changes": self.changes,
            "summary": dict(self.summary),
            "doc_object": doc if not preview_only else None
        }


# =============================================================================
# BATCH PROCESSOR
# =============================================================================

class BatchProcessor:
    """Processes multiple documents in batch."""

    def __init__(self, config: Config):
        self.config = config
        self.terms_dict = TermsDictionary(config.terms_dict_path)
        self.replacements_config = ReplacementsConfig(config.replacements_path)
        self.results = {}
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    def process_all(self, preview_only: bool = True) -> dict:
        """Process all documents."""
        documents = self.config.get_documents()
        replacements = self.replacements_config.get_ordered()

        if not documents:
            print("No documents found to process.")
            return {}

        if not replacements:
            print("No replacements defined in replacements.json")
            return {}

        print(f"\n{'PREVIEW MODE' if preview_only else 'APPLYING CHANGES'}")
        print(f"Documents to process: {len(documents)}")
        print(f"Replacements defined: {len(replacements)}")
        print("-" * 50)

        all_results = {
            "timestamp": self.timestamp,
            "mode": "preview" if preview_only else "apply",
            "documents": {},
            "total_changes": 0,
            "summary": defaultdict(int)
        }

        for doc_path in documents:
            print(f"\nProcessing: {doc_path.name}...")

            processor = DocumentProcessor(doc_path, self.config)
            result = processor.process(replacements, preview_only)

            all_results["documents"][doc_path.name] = {
                "change_count": result["change_count"],
                "changes": result["changes"],
                "summary": result["summary"]
            }
            all_results["total_changes"] += result["change_count"]

            for key, count in result["summary"].items():
                all_results["summary"][key] += count

            # Save document if not preview
            if not preview_only and result["doc_object"]:
                self.config.ensure_dirs()

                # Backup original
                backup_path = self.config.backup_dir / f"{doc_path.stem}_BACKUP_{self.timestamp}{doc_path.suffix}"
                shutil.copy2(doc_path, backup_path)
                print(f"  Backed up to: {backup_path.name}")

                # Save updated
                output_path = self.config.output_dir / f"{doc_path.stem}_UPDATED_{self.timestamp}{doc_path.suffix}"
                result["doc_object"].save(str(output_path))
                print(f"  Saved to: {output_path.name}")

            print(f"  Changes: {result['change_count']}")

        all_results["summary"] = dict(all_results["summary"])
        self.results = all_results
        return all_results

    def generate_report(self, output_path: Path = None) -> str:
        """Generate a text report of the results."""
        lines = []
        lines.append("=" * 80)
        lines.append(f"FEDRAMP PACKAGE {'PREVIEW' if self.results.get('mode') == 'preview' else 'UPDATE'} REPORT")
        lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("=" * 80)

        lines.append(f"\nTotal documents processed: {len(self.results.get('documents', {}))}")
        lines.append(f"Total changes: {self.results.get('total_changes', 0)}")

        # Summary by replacement
        lines.append("\n" + "-" * 40)
        lines.append("CHANGES BY REPLACEMENT")
        lines.append("-" * 40)

        for change, count in sorted(self.results.get("summary", {}).items(), key=lambda x: -x[1]):
            lines.append(f"  [{count:4d}x] {change}")

        # Per-document details
        lines.append("\n" + "-" * 40)
        lines.append("PER-DOCUMENT DETAILS")
        lines.append("-" * 40)

        for doc_name, doc_data in self.results.get("documents", {}).items():
            lines.append(f"\n{doc_name}:")
            lines.append(f"  Total changes: {doc_data['change_count']}")

            if doc_data.get("summary"):
                for change, count in sorted(doc_data["summary"].items(), key=lambda x: -x[1]):
                    lines.append(f"    [{count:3d}x] {change}")

        lines.append("\n" + "=" * 80)
        lines.append("END OF REPORT")
        lines.append("=" * 80)

        report = "\n".join(lines)

        if output_path:
            with open(output_path, 'w') as f:
                f.write(report)

        return report

    def generate_detailed_report(self, output_path: Path = None) -> str:
        """Generate a detailed text report showing each change with context."""
        lines = []
        lines.append("=" * 80)
        lines.append(f"FEDRAMP PACKAGE DETAILED {'PREVIEW' if self.results.get('mode') == 'preview' else 'CHANGES'} REPORT")
        lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("=" * 80)

        lines.append(f"\nTotal documents processed: {len(self.results.get('documents', {}))}")
        lines.append(f"Total changes: {self.results.get('total_changes', 0)}")

        # Group changes by replacement type for easier review
        for doc_name, doc_data in self.results.get("documents", {}).items():
            lines.append("\n" + "=" * 80)
            lines.append(f"DOCUMENT: {doc_name}")
            lines.append(f"Total changes: {doc_data['change_count']}")
            lines.append("=" * 80)

            # Group changes by old_text -> new_text
            changes_by_type = defaultdict(list)
            for change in doc_data.get("changes", []):
                key = f"{change['old_text']} -> {change['new_text']}"
                changes_by_type[key].append(change)

            # Sort by count (most frequent first)
            sorted_types = sorted(changes_by_type.items(), key=lambda x: -len(x[1]))

            for change_type, changes in sorted_types:
                lines.append("\n" + "-" * 60)
                lines.append(f"[{len(changes)}x] {change_type}")
                lines.append("-" * 60)

                # Show up to 5 examples with context
                for i, change in enumerate(changes[:5]):
                    lines.append(f"\n  Location: {change['location']}")
                    context = change.get('context_before', '')
                    if context:
                        # Highlight the term in context
                        lines.append(f"  Context:  \"{context}\"")
                    if change.get('is_deletion'):
                        lines.append(f"  Action:   DELETE")

                if len(changes) > 5:
                    lines.append(f"\n  ... and {len(changes) - 5} more instances")

        lines.append("\n" + "=" * 80)
        lines.append("END OF DETAILED REPORT")
        lines.append("=" * 80)

        report = "\n".join(lines)

        if output_path:
            with open(output_path, 'w') as f:
                f.write(report)

        return report

    def export_to_excel(self, output_path: Path) -> bool:
        """Export results to Excel spreadsheet."""
        if not PANDAS_AVAILABLE or not EXCEL_AVAILABLE:
            print("Error: pandas and openpyxl required for Excel export")
            print("Install with: pip install pandas openpyxl")
            return False

        # Create DataFrames
        # 1. Summary sheet
        summary_data = []
        for change, count in self.results.get("summary", {}).items():
            parts = change.split(" -> ")
            summary_data.append({
                "Original Term": parts[0] if len(parts) > 0 else "",
                "Replacement": parts[1] if len(parts) > 1 else "",
                "Total Count": count
            })
        summary_df = pd.DataFrame(summary_data)

        # 2. Document summary sheet
        doc_summary_data = []
        for doc_name, doc_data in self.results.get("documents", {}).items():
            doc_summary_data.append({
                "Document": doc_name,
                "Total Changes": doc_data["change_count"]
            })
        doc_summary_df = pd.DataFrame(doc_summary_data)

        # 3. Detailed changes sheet
        changes_data = []
        for doc_name, doc_data in self.results.get("documents", {}).items():
            for change in doc_data.get("changes", []):
                changes_data.append({
                    "Document": doc_name,
                    "Location": change["location"],
                    "Original Term": change["old_text"],
                    "Replacement": change["new_text"],
                    "Is Deletion": change.get("is_deletion", False),
                    "Context": change.get("context_before", "")[:100]
                })
        changes_df = pd.DataFrame(changes_data)

        # Write to Excel
        with pd.ExcelWriter(str(output_path), engine='openpyxl') as writer:
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
            doc_summary_df.to_excel(writer, sheet_name='Documents', index=False)
            changes_df.to_excel(writer, sheet_name='All Changes', index=False)

        return True

    def export_to_csv(self, output_dir: Path) -> bool:
        """Export results to CSV files."""
        if not PANDAS_AVAILABLE:
            print("Error: pandas required for CSV export")
            return False

        output_dir.mkdir(exist_ok=True)

        # Summary
        summary_data = []
        for change, count in self.results.get("summary", {}).items():
            parts = change.split(" -> ")
            summary_data.append({
                "Original Term": parts[0] if len(parts) > 0 else "",
                "Replacement": parts[1] if len(parts) > 1 else "",
                "Total Count": count
            })
        pd.DataFrame(summary_data).to_csv(output_dir / "summary.csv", index=False)

        # All changes
        changes_data = []
        for doc_name, doc_data in self.results.get("documents", {}).items():
            for change in doc_data.get("changes", []):
                changes_data.append({
                    "Document": doc_name,
                    "Location": change["location"],
                    "Original Term": change["old_text"],
                    "Replacement": change["new_text"],
                    "Is Deletion": change.get("is_deletion", False),
                    "Context": change.get("context_before", "")
                })
        pd.DataFrame(changes_data).to_csv(output_dir / "all_changes.csv", index=False)

        return True


# =============================================================================
# ANALYZER (for discovery)
# =============================================================================

class PackageAnalyzer:
    """Analyzes documents for term discovery."""

    def __init__(self, config: Config):
        self.config = config
        self.terms_dict = TermsDictionary(config.terms_dict_path)
        self.results = {}

    def analyze(self) -> dict:
        """Analyze all documents."""
        documents = self.config.get_documents()
        known_terms = self.terms_dict.get_all_known_terms()

        print(f"\nAnalyzing {len(documents)} documents...")

        results = {
            "timestamp": datetime.now().isoformat(),
            "documents_analyzed": len(documents),
            "known_terms": defaultdict(lambda: {"documents": [], "total_count": 0}),
            "per_document": {}
        }

        for doc_path in documents:
            print(f"  {doc_path.name}...")

            try:
                doc = Document(str(doc_path))
                processor = DocumentProcessor(doc_path, self.config)
                full_text = processor.extract_text(doc)

                doc_terms = {}
                for term, info in known_terms.items():
                    pattern = re.compile(rf'\b{re.escape(term)}\b')
                    matches = pattern.findall(full_text)
                    if matches:
                        doc_terms[term] = len(matches)
                        results["known_terms"][term]["documents"].append({
                            "file": doc_path.name,
                            "count": len(matches)
                        })
                        results["known_terms"][term]["total_count"] += len(matches)
                        results["known_terms"][term]["type"] = info.get("type")

                results["per_document"][doc_path.name] = {
                    "terms_found": len(doc_terms),
                    "terms": doc_terms
                }

            except Exception as e:
                print(f"    Error: {e}")
                results["per_document"][doc_path.name] = {"error": str(e)}

        results["known_terms"] = dict(results["known_terms"])
        self.results = results
        return results

    def export_to_excel(self, output_path: Path) -> bool:
        """Export analysis to Excel."""
        if not PANDAS_AVAILABLE or not EXCEL_AVAILABLE:
            return False

        # Terms summary
        terms_data = []
        for term, data in self.results.get("known_terms", {}).items():
            terms_data.append({
                "Term": term,
                "Type": data.get("type", "unknown"),
                "Total Count": data["total_count"],
                "Documents": len(data["documents"]),
                "Document List": ", ".join([d["file"] for d in data["documents"]])
            })
        terms_df = pd.DataFrame(terms_data)

        # Cross-document matrix
        all_terms = list(self.results.get("known_terms", {}).keys())
        all_docs = list(self.results.get("per_document", {}).keys())

        matrix_data = []
        for doc in all_docs:
            row = {"Document": doc}
            doc_terms = self.results["per_document"].get(doc, {}).get("terms", {})
            for term in all_terms:
                row[term] = doc_terms.get(term, 0)
            matrix_data.append(row)
        matrix_df = pd.DataFrame(matrix_data)

        with pd.ExcelWriter(str(output_path), engine='openpyxl') as writer:
            terms_df.to_excel(writer, sheet_name='Terms Summary', index=False)
            matrix_df.to_excel(writer, sheet_name='Cross-Document Matrix', index=False)

        return True


# =============================================================================
# FONT ANALYSIS AND STANDARDIZATION
# =============================================================================

def format_font_combo(font_name, size, bold, italic) -> str:
    """Format a font combination for display."""
    parts = []
    parts.append(font_name if font_name else "[inherited]")
    parts.append(f"{size}pt" if size else "[inherited]")
    if bold is True:
        parts.append("Bold")
    if italic is True:
        parts.append("Italic")
    return ", ".join(parts)


class FontAnalyzer:
    """Analyzes font usage across documents."""

    def __init__(self, config: Config):
        self.config = config
        self.results = {}
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    def analyze_document(self, doc_path: Path) -> dict:
        """Analyze all fonts in a single document."""
        doc = Document(str(doc_path))

        # Track font combinations: (name, size, bold, italic) -> {count, samples, locations}
        font_combinations = defaultdict(lambda: {
            "count": 0,
            "samples": [],
            "locations": []
        })

        def process_paragraph(para, location: str):
            for run in para.runs:
                if run.text.strip():
                    font = run.font
                    key = (
                        font.name,
                        font.size.pt if font.size else None,
                        font.bold,
                        font.italic
                    )
                    combo = font_combinations[key]
                    combo["count"] += 1
                    if len(combo["samples"]) < 3:
                        combo["samples"].append(run.text.strip()[:60])
                    if len(combo["locations"]) < 5:
                        combo["locations"].append(location)

        # Process all document sections
        for i, para in enumerate(doc.paragraphs):
            process_paragraph(para, f"Body, Paragraph {i+1}")

        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        process_paragraph(para, f"Table {ti+1}, Row {ri+1}, Col {ci+1}")

        for si, section in enumerate(doc.sections):
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        process_paragraph(para, f"Header (Section {si+1})")
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        process_paragraph(para, f"Footer (Section {si+1})")

        # Convert defaultdict keys to serializable format
        result_combos = {}
        for key, data in font_combinations.items():
            str_key = f"{key[0]}|{key[1]}|{key[2]}|{key[3]}"
            result_combos[str_key] = {
                "font_name": key[0],
                "font_size_pt": key[1],
                "bold": key[2],
                "italic": key[3],
                **data
            }

        return {
            "document": doc_path.name,
            "font_combinations": result_combos,
            "total_runs": sum(c["count"] for c in result_combos.values())
        }

    def analyze_all(self) -> dict:
        """Analyze all documents in the package."""
        documents = self.config.get_drafts() or self.config.get_documents()

        if not documents:
            return {"error": "No documents found"}

        all_combinations = defaultdict(lambda: {
            "count": 0,
            "samples": [],
            "locations": [],
            "documents": []
        })

        doc_results = {}

        for doc_path in documents:
            print(f"  Analyzing: {doc_path.name}...")
            result = self.analyze_document(doc_path)
            doc_results[doc_path.name] = result

            # Aggregate combinations across documents
            for combo_key, combo_data in result["font_combinations"].items():
                agg = all_combinations[combo_key]
                agg["count"] += combo_data["count"]
                agg["font_name"] = combo_data["font_name"]
                agg["font_size_pt"] = combo_data["font_size_pt"]
                agg["bold"] = combo_data["bold"]
                agg["italic"] = combo_data["italic"]
                if len(agg["samples"]) < 3:
                    agg["samples"].extend(combo_data["samples"][:3 - len(agg["samples"])])
                if doc_path.name not in agg["documents"]:
                    agg["documents"].append(doc_path.name)

        self.results = {
            "timestamp": datetime.now().isoformat(),
            "documents_analyzed": len(documents),
            "documents": doc_results,
            "all_combinations": dict(all_combinations),
            "total_runs": sum(d["total_runs"] for d in doc_results.values())
        }

        return self.results

    def generate_report(self, output_path: Path = None) -> str:
        """Generate a text report of font analysis."""
        lines = []
        lines.append("=" * 80)
        lines.append("FONT USAGE REPORT")
        lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("=" * 80)

        lines.append(f"\nDocuments analyzed: {self.results.get('documents_analyzed', 0)}")
        lines.append(f"Total text runs: {self.results.get('total_runs', 0)}")
        lines.append(f"Unique font combinations: {len(self.results.get('all_combinations', {}))}")

        lines.append("\n" + "-" * 40)
        lines.append("FONT COMBINATIONS BY FREQUENCY")
        lines.append("-" * 40)

        sorted_combos = sorted(
            self.results.get("all_combinations", {}).items(),
            key=lambda x: -x[1]["count"]
        )

        total_runs = self.results.get("total_runs", 1)
        for combo_key, data in sorted_combos:
            desc = format_font_combo(
                data["font_name"], data["font_size_pt"],
                data["bold"], data["italic"]
            )
            pct = (data["count"] / total_runs) * 100 if total_runs > 0 else 0

            lines.append(f"\n[{data['count']:6d}x] ({pct:5.1f}%) {desc}")
            if data["samples"]:
                sample = data["samples"][0][:50]
                lines.append(f"           Sample: \"{sample}...\"")
            if data["documents"]:
                docs = ", ".join(data["documents"][:3])
                lines.append(f"           Documents: {docs}")

        # Per-document breakdown
        lines.append("\n" + "-" * 40)
        lines.append("PER-DOCUMENT SUMMARY")
        lines.append("-" * 40)

        for doc_name, doc_data in self.results.get("documents", {}).items():
            lines.append(f"\n{doc_name}:")
            lines.append(f"  Total runs: {doc_data['total_runs']}")
            lines.append(f"  Font combinations: {len(doc_data['font_combinations'])}")

        lines.append("\n" + "=" * 80)
        lines.append("END OF REPORT")
        lines.append("=" * 80)

        report = "\n".join(lines)

        if output_path:
            with open(output_path, 'w') as f:
                f.write(report)

        return report

    def export_to_excel(self, output_path: Path) -> bool:
        """Export font analysis to Excel."""
        if not PANDAS_AVAILABLE or not EXCEL_AVAILABLE:
            return False

        # Summary sheet
        summary_data = []
        for combo_key, data in self.results.get("all_combinations", {}).items():
            summary_data.append({
                "Font Name": data["font_name"] or "[inherited]",
                "Size (pt)": data["font_size_pt"] or "[inherited]",
                "Bold": str(data["bold"]) if data["bold"] is not None else "[inherited]",
                "Italic": str(data["italic"]) if data["italic"] is not None else "[inherited]",
                "Count": data["count"],
                "Percentage": f"{(data['count'] / self.results.get('total_runs', 1)) * 100:.1f}%",
                "Sample Text": data["samples"][0][:60] if data["samples"] else ""
            })
        summary_df = pd.DataFrame(summary_data)
        summary_df = summary_df.sort_values("Count", ascending=False)

        # Per-document sheet
        doc_data = []
        for doc_name, doc_info in self.results.get("documents", {}).items():
            doc_data.append({
                "Document": doc_name,
                "Total Runs": doc_info["total_runs"],
                "Unique Font Combinations": len(doc_info["font_combinations"])
            })
        doc_df = pd.DataFrame(doc_data)

        with pd.ExcelWriter(str(output_path), engine='openpyxl') as writer:
            summary_df.to_excel(writer, sheet_name='Font Summary', index=False)
            doc_df.to_excel(writer, sheet_name='Per Document', index=False)

        return True


class FontStandardizer:
    """Standardizes fonts across documents."""

    # Fonts to preserve (symbols, checkboxes)
    PRESERVE_FONTS = ["Segoe UI Symbol", "Wingdings", "Symbol", "Webdings"]

    def __init__(self, config: Config, target_font: str = "Arial",
                 target_size: float = 11.0):
        self.config = config
        self.target_font = target_font
        self.target_size = target_size
        self.changes = []
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    def should_standardize(self, run) -> bool:
        """Determine if this run should be standardized."""
        # Don't change symbol fonts
        if run.font.name in self.PRESERVE_FONTS:
            return False

        # Check for special characters only (checkboxes, etc.)
        text = run.text.strip()
        if text and all(ord(c) > 8000 for c in text):
            return False

        # Skip empty runs
        if not text:
            return False

        return True

    def standardize_run(self, run, location: str, preview_only: bool = True) -> dict:
        """Standardize a single run's font."""
        if not self.should_standardize(run):
            return None

        from docx.shared import Pt

        font = run.font
        current_name = font.name
        current_size = font.size.pt if font.size else None

        # Check if change is needed (only explicit fonts, or if we want to set all)
        # For now, we'll set explicit fonts on everything that's not a symbol
        change = {
            "location": location,
            "text_sample": run.text[:40],
            "before": {
                "font_name": current_name,
                "font_size_pt": current_size
            },
            "after": {
                "font_name": self.target_font,
                "font_size_pt": self.target_size
            }
        }

        if not preview_only:
            font.name = self.target_font
            font.size = Pt(self.target_size)

        return change

    def standardize_document(self, doc_path: Path, preview_only: bool = True) -> dict:
        """Standardize all fonts in a document."""
        doc = Document(str(doc_path))
        changes = []

        def process_paragraph(para, location: str):
            for run in para.runs:
                change = self.standardize_run(run, location, preview_only)
                if change:
                    changes.append(change)

        # Process all sections
        for i, para in enumerate(doc.paragraphs):
            process_paragraph(para, f"Body, Paragraph {i+1}")

        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        process_paragraph(para, f"Table {ti+1}, Row {ri+1}, Col {ci+1}")

        for si, section in enumerate(doc.sections):
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        process_paragraph(para, f"Header (Section {si+1})")
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        process_paragraph(para, f"Footer (Section {si+1})")

        # Save if not preview
        if not preview_only and changes:
            doc.save(str(doc_path))

        return {
            "document": doc_path.name,
            "changes": changes,
            "change_count": len(changes)
        }

    def standardize_drafts(self, preview_only: bool = True) -> dict:
        """Standardize fonts in all draft documents."""
        drafts = self.config.get_drafts()

        results = {
            "timestamp": datetime.now().isoformat(),
            "mode": "preview" if preview_only else "apply",
            "target_font": self.target_font,
            "target_size_pt": self.target_size,
            "documents": {},
            "total_changes": 0
        }

        for draft_path in drafts:
            print(f"  Processing: {draft_path.name}...")
            doc_result = self.standardize_document(draft_path, preview_only)
            results["documents"][draft_path.name] = doc_result
            results["total_changes"] += doc_result["change_count"]
            print(f"    Changes: {doc_result['change_count']}")

        return results


# =============================================================================
# CLI COMMANDS
# =============================================================================

def cmd_status(config: Config):
    """Show current configuration status."""
    print("\n" + "=" * 60)
    print("PACKAGE MANAGER STATUS")
    print("=" * 60)

    print(f"\nBase directory: {config.base_dir.absolute()}")

    # Directory structure
    print("\n" + "-" * 40)
    print("DIRECTORY STRUCTURE")
    print("-" * 40)

    originals = config.get_originals()
    drafts = config.get_drafts()

    orig_status = f"{len(originals)} documents" if config.originals_dir.exists() else "NOT CREATED"
    draft_status = f"{len(drafts)} documents" if config.drafts_dir.exists() else "NOT CREATED"
    output_status = "EXISTS" if config.output_dir.exists() else "NOT CREATED"
    backup_status = "EXISTS" if config.backup_dir.exists() else "NOT CREATED"

    print(f"\n  originals/  : {orig_status}")
    print(f"  drafts/     : {draft_status}")
    print(f"  output/     : {output_status}")
    print(f"  backups/    : {backup_status}")

    if originals:
        print(f"\n  Original documents:")
        for doc in originals[:5]:
            print(f"    - {doc.name}")
        if len(originals) > 5:
            print(f"    ... and {len(originals) - 5} more")

    if drafts:
        print(f"\n  Draft documents:")
        for doc in drafts[:5]:
            print(f"    - {doc.name}")
        if len(drafts) > 5:
            print(f"    ... and {len(drafts) - 5} more")

    # Check for documents in base directory (legacy)
    base_docs = []
    for pattern in config.include_patterns:
        base_docs.extend(config.base_dir.glob(pattern))
    base_docs = config._filter_docs(base_docs)
    if base_docs and not originals:
        print(f"\n  NOTE: Found {len(base_docs)} documents in base directory.")
        print(f"        Run 'apply' to move them to originals/ and start workflow.")

    # Terms dictionary
    print("\n" + "-" * 40)
    print("CONFIGURATION FILES")
    print("-" * 40)

    if config.terms_dict_path.exists():
        terms_dict = TermsDictionary(config.terms_dict_path)
        known = terms_dict.get_all_known_terms()
        print(f"\n  Terms dictionary: {config.terms_dict_path.name}")
        print(f"    Known terms: {len(known)}")
    else:
        print(f"\n  Terms dictionary: NOT FOUND ({config.terms_dict_path.name})")
        print(f"    Copy from examples/terms_dictionary.example.json")

    # Replacements
    if config.replacements_path.exists():
        replacements = ReplacementsConfig(config.replacements_path)
        print(f"\n  Replacements file: {config.replacements_path.name}")
        print(f"    Active replacements: {len(replacements.replacements)}")
        for old, new in list(replacements.replacements.items())[:5]:
            display_new = new if new else "[DELETE]"
            print(f"      '{old}' -> '{display_new}'")
        if len(replacements.replacements) > 5:
            print(f"      ... and {len(replacements.replacements) - 5} more")
    else:
        print(f"\n  Replacements file: NOT FOUND ({config.replacements_path.name})")
        print(f"    Copy from examples/replacements.example.json")

    # Workflow status
    print("\n" + "-" * 40)
    print("WORKFLOW STATUS")
    print("-" * 40)

    if not originals and not base_docs:
        print("\n  Step 1: Place documents in 'originals/' directory")
    elif originals and not drafts:
        print("\n  Ready for: analyze -> preview -> apply")
    elif drafts:
        print("\n  Drafts exist. Ready for: verify")
        print("  Or run 'apply' again to regenerate drafts from originals.")


def cmd_analyze(config: Config, export_format: str = None):
    """Analyze documents for terms."""
    print("\n" + "=" * 60)
    print("ANALYZING PACKAGE DOCUMENTS")
    print("=" * 60)

    analyzer = PackageAnalyzer(config)
    results = analyzer.analyze()

    # Print summary
    print("\n" + "-" * 40)
    print("ANALYSIS SUMMARY")
    print("-" * 40)

    print(f"\nDocuments analyzed: {results['documents_analyzed']}")
    print(f"Known terms found: {len(results['known_terms'])}")

    # Top terms
    sorted_terms = sorted(
        results["known_terms"].items(),
        key=lambda x: -x[1]["total_count"]
    )[:15]

    print("\nTop 15 terms across all documents:")
    for term, data in sorted_terms:
        doc_count = len(data["documents"])
        print(f"  [{data['total_count']:4d}x in {doc_count} doc(s)] {term}")

    # Save results
    config.ensure_dirs()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    json_path = config.output_dir / f"analysis_{timestamp}.json"
    with open(json_path, 'w') as f:
        json.dump(results, f, indent=2, default=str)
    print(f"\nJSON results: {json_path}")

    if export_format == "excel":
        excel_path = config.output_dir / f"analysis_{timestamp}.xlsx"
        if analyzer.export_to_excel(excel_path):
            print(f"Excel export: {excel_path}")


def cmd_preview(config: Config, export_format: str = None):
    """Preview changes across all documents."""
    print("\n" + "=" * 60)
    print("PREVIEWING CHANGES")
    print("=" * 60)

    processor = BatchProcessor(config)
    results = processor.process_all(preview_only=True)

    if not results:
        return

    # Generate reports
    config.ensure_dirs()

    # Summary report
    report_path = config.output_dir / f"preview_{processor.timestamp}.txt"
    processor.generate_report(report_path)
    print(f"\nSummary report: {report_path}")

    # Detailed report with context
    detailed_path = config.output_dir / f"preview_{processor.timestamp}_detailed.txt"
    processor.generate_detailed_report(detailed_path)
    print(f"Detailed report (with context): {detailed_path}")

    # JSON
    json_path = config.output_dir / f"preview_{processor.timestamp}.json"
    with open(json_path, 'w') as f:
        json.dump(results, f, indent=2, default=str)
    print(f"JSON results: {json_path}")

    # Excel export
    if export_format == "excel":
        excel_path = config.output_dir / f"preview_{processor.timestamp}.xlsx"
        if processor.export_to_excel(excel_path):
            print(f"Excel export: {excel_path}")
    elif export_format == "csv":
        csv_dir = config.output_dir / f"preview_{processor.timestamp}_csv"
        if processor.export_to_csv(csv_dir):
            print(f"CSV export: {csv_dir}")

    # Summary
    print("\n" + "-" * 40)
    print("PREVIEW SUMMARY")
    print("-" * 40)
    print(f"Total changes that would be made: {results['total_changes']}")
    print("\nRun 'python package_manager.py apply' to apply these changes.")


def cmd_apply(config: Config, export_format: str = None):
    """Apply changes to all documents using the originals -> drafts workflow."""
    print("\n" + "=" * 60)
    print("APPLYING CHANGES")
    print("=" * 60)

    config.ensure_dirs()

    # Check for originals
    originals = config.get_originals()
    replacements = ReplacementsConfig(config.replacements_path)

    if not originals:
        # Check for documents in base directory
        base_docs = []
        for pattern in config.include_patterns:
            base_docs.extend(config.base_dir.glob(pattern))
        base_docs = config._filter_docs(base_docs)

        if base_docs:
            print(f"\nNo documents found in 'originals/' directory.")
            print(f"Found {len(base_docs)} documents in base directory.")
            print("\nWould you like to move them to 'originals/' to use the")
            print("recommended workflow? (originals are preserved, changes go to drafts)")
            response = input("\nMove documents to originals/? [Y/n]: ").strip().lower()

            if response != 'n':
                for doc in base_docs:
                    dest = config.originals_dir / doc.name
                    shutil.move(str(doc), str(dest))
                    print(f"  Moved: {doc.name}")
                originals = config.get_originals()
            else:
                print("\nPlease place source documents in the 'originals/' directory.")
                return
        else:
            print("\nNo documents found. Place source documents in 'originals/' directory.")
            return

    print(f"\nWorkflow:")
    print(f"  1. Copy originals from: originals/")
    print(f"  2. Apply changes to:    drafts/")
    print(f"  3. Original files:      PRESERVED (never modified)")
    print(f"\nDocuments to process: {len(originals)}")
    print(f"Replacements defined: {len(replacements.replacements)}")

    response = input("\nProceed? [y/N]: ").strip().lower()

    if response != 'y':
        print("Cancelled.")
        return

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Copy originals to drafts (or backup existing drafts first)
    print("\n" + "-" * 40)
    print("PREPARING DRAFTS")
    print("-" * 40)

    for orig_doc in originals:
        draft_path = config.drafts_dir / orig_doc.name

        # If draft exists, back it up
        if draft_path.exists():
            backup_path = config.backup_dir / f"{orig_doc.stem}_DRAFT_BACKUP_{timestamp}{orig_doc.suffix}"
            shutil.copy2(draft_path, backup_path)
            print(f"  Backed up existing draft: {draft_path.name}")

        # Copy original to draft
        shutil.copy2(orig_doc, draft_path)
        print(f"  Created draft: {draft_path.name}")

    # Now process the drafts
    print("\n" + "-" * 40)
    print("APPLYING REPLACEMENTS TO DRAFTS")
    print("-" * 40)

    ordered_replacements = replacements.get_ordered()
    all_results = {
        "timestamp": timestamp,
        "mode": "apply",
        "documents": {},
        "total_changes": 0,
        "summary": defaultdict(int)
    }

    drafts = config.get_drafts()
    for draft_path in drafts:
        print(f"\nProcessing: {draft_path.name}...")

        processor = DocumentProcessor(draft_path, config)
        result = processor.process(ordered_replacements, preview_only=False)

        all_results["documents"][draft_path.name] = {
            "change_count": result["change_count"],
            "changes": result["changes"],
            "summary": result["summary"]
        }
        all_results["total_changes"] += result["change_count"]

        for key, count in result["summary"].items():
            all_results["summary"][key] += count

        # Save the modified document back to drafts
        if result["doc_object"]:
            result["doc_object"].save(str(draft_path))
            print(f"  Changes saved: {result['change_count']}")

    all_results["summary"] = dict(all_results["summary"])

    # Generate reports
    print("\n" + "-" * 40)
    print("GENERATING REPORTS")
    print("-" * 40)

    # Create a BatchProcessor just for report generation
    batch = BatchProcessor(config)
    batch.results = all_results
    batch.timestamp = timestamp

    report_path = config.output_dir / f"applied_{timestamp}.txt"
    batch.generate_report(report_path)
    print(f"Text report: {report_path}")

    json_path = config.output_dir / f"applied_{timestamp}.json"
    with open(json_path, 'w') as f:
        json.dump(all_results, f, indent=2, default=str)
    print(f"JSON results: {json_path}")

    if export_format == "excel":
        excel_path = config.output_dir / f"applied_{timestamp}.xlsx"
        if batch.export_to_excel(excel_path):
            print(f"Excel export: {excel_path}")

    print("\n" + "=" * 60)
    print("COMPLETE")
    print("=" * 60)
    print(f"Total changes applied: {all_results['total_changes']}")
    print(f"\nOriginals preserved in: {config.originals_dir}")
    print(f"Updated drafts in:      {config.drafts_dir}")
    print(f"Reports saved to:       {config.output_dir}")
    print(f"\nNext step: Run 'python package_manager.py verify' to validate the drafts.")


def cmd_export(config: Config, export_format: str):
    """Export current analysis to spreadsheet."""
    print("\n" + "=" * 60)
    print("EXPORTING DATA")
    print("=" * 60)

    # Run analysis first
    analyzer = PackageAnalyzer(config)
    results = analyzer.analyze()

    config.ensure_dirs()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if export_format == "excel":
        excel_path = config.output_dir / f"package_export_{timestamp}.xlsx"
        if analyzer.export_to_excel(excel_path):
            print(f"\nExcel file created: {excel_path}")
            print("\nSheets included:")
            print("  - Terms Summary: All known terms with counts")
            print("  - Cross-Document Matrix: Terms by document")
        else:
            print("Failed to create Excel export.")
    else:
        print(f"Unknown format: {export_format}")


def cmd_verify(config: Config, export_format: str = None):
    """Verify drafts by re-analyzing them for completeness and remaining terms."""
    print("\n" + "=" * 60)
    print("VERIFYING DRAFT DOCUMENTS")
    print("=" * 60)

    drafts = config.get_drafts()
    originals = config.get_originals()

    if not drafts:
        print("\nNo draft documents found in 'drafts/' directory.")
        print("Run 'python package_manager.py apply' first to create drafts.")
        return

    print(f"\nDrafts to verify: {len(drafts)}")
    print(f"Originals for comparison: {len(originals)}")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    terms_dict = TermsDictionary(config.terms_dict_path)
    known_terms = terms_dict.get_all_known_terms()
    replacements = ReplacementsConfig(config.replacements_path)

    verification_results = {
        "timestamp": timestamp,
        "drafts_analyzed": len(drafts),
        "documents": {},
        "remaining_terms": defaultdict(list),
        "replacement_verification": {},
        "summary": {
            "total_original_terms_remaining": 0,
            "all_replacements_applied": True,
            "documents_with_issues": []
        }
    }

    # Get list of terms that should have been replaced
    replaced_terms = set(replacements.replacements.keys())

    print("\n" + "-" * 40)
    print("ANALYZING DRAFTS")
    print("-" * 40)

    for draft_path in drafts:
        print(f"\n  Verifying: {draft_path.name}...")

        try:
            doc = Document(str(draft_path))
            processor = DocumentProcessor(draft_path, config)
            full_text = processor.extract_text(doc)

            doc_results = {
                "terms_found": {},
                "original_terms_remaining": [],
                "issues": []
            }

            # Check for known terms still present
            for term, info in known_terms.items():
                pattern = re.compile(rf'\b{re.escape(term)}\b')
                matches = pattern.findall(full_text)
                if matches:
                    doc_results["terms_found"][term] = {
                        "count": len(matches),
                        "type": info.get("type")
                    }

                    # Was this term supposed to be replaced?
                    if term in replaced_terms:
                        doc_results["original_terms_remaining"].append({
                            "term": term,
                            "count": len(matches),
                            "expected_replacement": replacements.replacements.get(term)
                        })
                        verification_results["remaining_terms"][term].append({
                            "document": draft_path.name,
                            "count": len(matches)
                        })
                        verification_results["summary"]["total_original_terms_remaining"] += len(matches)
                        verification_results["summary"]["all_replacements_applied"] = False

            if doc_results["original_terms_remaining"]:
                doc_results["issues"].append(
                    f"Found {len(doc_results['original_terms_remaining'])} terms that should have been replaced"
                )
                if draft_path.name not in verification_results["summary"]["documents_with_issues"]:
                    verification_results["summary"]["documents_with_issues"].append(draft_path.name)

            verification_results["documents"][draft_path.name] = doc_results

            # Print summary for this doc
            remaining_count = len(doc_results["original_terms_remaining"])
            if remaining_count > 0:
                print(f"    WARNING: {remaining_count} original terms still present")
            else:
                print(f"    OK: All replacements applied successfully")

        except Exception as e:
            print(f"    ERROR: {e}")
            verification_results["documents"][draft_path.name] = {"error": str(e)}

    verification_results["remaining_terms"] = dict(verification_results["remaining_terms"])

    # Generate verification report
    print("\n" + "-" * 40)
    print("VERIFICATION SUMMARY")
    print("-" * 40)

    if verification_results["summary"]["all_replacements_applied"]:
        print("\n  SUCCESS: All replacements have been applied correctly!")
    else:
        print(f"\n  WARNING: {verification_results['summary']['total_original_terms_remaining']} "
              f"original terms still found in drafts")
        print("\n  Terms that should have been replaced but remain:")
        for term, occurrences in verification_results["remaining_terms"].items():
            total = sum(o["count"] for o in occurrences)
            print(f"    [{total:3d}x] {term}")

    # Compare originals vs drafts term counts
    print("\n" + "-" * 40)
    print("ORIGINALS VS DRAFTS COMPARISON")
    print("-" * 40)

    if originals:
        # Analyze originals for comparison
        original_terms = defaultdict(int)
        for orig_path in originals:
            try:
                doc = Document(str(orig_path))
                processor = DocumentProcessor(orig_path, config)
                full_text = processor.extract_text(doc)
                for term in replaced_terms:
                    pattern = re.compile(rf'\b{re.escape(term)}\b')
                    matches = pattern.findall(full_text)
                    original_terms[term] += len(matches)
            except Exception:
                pass

        draft_terms = defaultdict(int)
        for draft_path in drafts:
            for term, occurrences in verification_results["remaining_terms"].items():
                for occ in occurrences:
                    if occ["document"] == draft_path.name:
                        draft_terms[term] += occ["count"]

        print("\n  Replacement term changes:")
        print(f"  {'Term':<40} {'Originals':>10} {'Drafts':>10} {'Status':>12}")
        print("  " + "-" * 74)

        for term in sorted(replaced_terms):
            orig_count = original_terms.get(term, 0)
            draft_count = draft_terms.get(term, 0)
            if orig_count > 0 or draft_count > 0:
                if draft_count == 0:
                    status = "REPLACED"
                elif draft_count < orig_count:
                    status = "PARTIAL"
                else:
                    status = "UNCHANGED"
                print(f"  {term:<40} {orig_count:>10} {draft_count:>10} {status:>12}")

    # Check for repeated words (e.g., "Knox Knox")
    print("\n" + "-" * 40)
    print("CHECKING FOR REPEATED WORDS")
    print("-" * 40)

    repeated_word_issues = []
    for draft_path in drafts:
        result = check_document_for_repeated_words(draft_path)
        if result["total_issues"] > 0:
            repeated_word_issues.extend([
                {**issue, "document": result["document"]}
                for issue in result["issues"]
            ])

    if repeated_word_issues:
        print(f"\n  WARNING: Found {len(repeated_word_issues)} repeated word issues:")
        for issue in repeated_word_issues[:20]:  # Show first 20
            print(f"\n    Document: {issue['document']}")
            print(f"    Location: {issue['location']}")
            print(f"    Issue:    '{issue['full_match']}'")
            print(f"    Context:  ...{issue['context']}...")
        if len(repeated_word_issues) > 20:
            print(f"\n    ... and {len(repeated_word_issues) - 20} more issues")

        # Add to verification results
        verification_results["repeated_words"] = repeated_word_issues
        verification_results["summary"]["repeated_word_count"] = len(repeated_word_issues)
    else:
        print("\n  OK: No repeated words found")
        verification_results["repeated_words"] = []
        verification_results["summary"]["repeated_word_count"] = 0

    # Save verification results
    config.ensure_dirs()
    json_path = config.output_dir / f"verification_{timestamp}.json"
    with open(json_path, 'w') as f:
        json.dump(verification_results, f, indent=2, default=str)
    print(f"\nVerification results: {json_path}")

    # Text report
    report_lines = []
    report_lines.append("=" * 80)
    report_lines.append("DRAFT VERIFICATION REPORT")
    report_lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report_lines.append("=" * 80)
    report_lines.append(f"\nDrafts verified: {len(drafts)}")
    report_lines.append(f"All replacements applied: {verification_results['summary']['all_replacements_applied']}")

    if not verification_results["summary"]["all_replacements_applied"]:
        report_lines.append(f"\nWARNING: The following original terms were NOT fully replaced:")
        for term, occurrences in verification_results["remaining_terms"].items():
            total = sum(o["count"] for o in occurrences)
            report_lines.append(f"  [{total:3d}x] {term}")
            for occ in occurrences:
                report_lines.append(f"         - {occ['document']}: {occ['count']} occurrences")

    # Add repeated words section
    if verification_results.get("repeated_words"):
        report_lines.append(f"\n" + "-" * 40)
        report_lines.append("REPEATED WORDS FOUND")
        report_lines.append("-" * 40)
        report_lines.append(f"\nTotal repeated word issues: {len(verification_results['repeated_words'])}")
        for issue in verification_results["repeated_words"]:
            report_lines.append(f"\n  Document: {issue['document']}")
            report_lines.append(f"  Location: {issue['location']}")
            report_lines.append(f"  Issue:    '{issue['full_match']}'")
            report_lines.append(f"  Context:  ...{issue['context']}...")

    report_lines.append("\n" + "=" * 80)
    report_lines.append("END OF REPORT")
    report_lines.append("=" * 80)

    report_path = config.output_dir / f"verification_{timestamp}.txt"
    with open(report_path, 'w') as f:
        f.write("\n".join(report_lines))
    print(f"Text report: {report_path}")

    if export_format == "excel" and PANDAS_AVAILABLE and EXCEL_AVAILABLE:
        # Create Excel verification report
        excel_data = []
        for term, occurrences in verification_results["remaining_terms"].items():
            for occ in occurrences:
                excel_data.append({
                    "Term": term,
                    "Document": occ["document"],
                    "Count": occ["count"],
                    "Expected Replacement": replacements.replacements.get(term, "")
                })

        if excel_data:
            df = pd.DataFrame(excel_data)
            excel_path = config.output_dir / f"verification_{timestamp}.xlsx"
            df.to_excel(str(excel_path), index=False)
            print(f"Excel report: {excel_path}")


def cmd_fonts(config: Config, export_format: str = None):
    """Analyze font usage across all draft documents."""
    print("\n" + "=" * 60)
    print("ANALYZING FONT USAGE")
    print("=" * 60)

    drafts = config.get_drafts()
    if not drafts:
        print("\nNo draft documents found in 'drafts/' directory.")
        print("Run 'python package_manager.py apply' first to create drafts.")
        return

    analyzer = FontAnalyzer(config)
    results = analyzer.analyze_all()

    if not results["documents"]:
        print("\nNo documents analyzed.")
        return

    # Generate reports
    config.ensure_dirs()
    timestamp = analyzer.timestamp

    # Text report
    report_path = config.output_dir / f"fonts_{timestamp}.txt"
    analyzer.generate_report(report_path)
    print(f"\nText report: {report_path}")

    # JSON
    json_path = config.output_dir / f"fonts_{timestamp}.json"
    with open(json_path, 'w') as f:
        json.dump(results, f, indent=2, default=str)
    print(f"JSON results: {json_path}")

    # Excel
    if export_format == "excel":
        excel_path = config.output_dir / f"fonts_{timestamp}.xlsx"
        if analyzer.export_to_excel(excel_path):
            print(f"Excel export: {excel_path}")

    print("\n" + "-" * 40)
    print("NEXT STEPS")
    print("-" * 40)
    print("\nTo standardize fonts to Arial 11pt:")
    print("  python package_manager.py standardize-fonts           # Preview changes")
    print("  python package_manager.py standardize-fonts --apply   # Apply changes")


def cmd_standardize_fonts(config: Config, apply: bool = False,
                          target_font: str = "Arial", target_size: float = 11.0):
    """Standardize fonts in draft documents."""
    mode_str = "APPLYING" if apply else "PREVIEWING"
    print("\n" + "=" * 60)
    print(f"{mode_str} FONT STANDARDIZATION")
    print("=" * 60)

    drafts = config.get_drafts()
    if not drafts:
        print("\nNo draft documents found in 'drafts/' directory.")
        print("Run 'python package_manager.py apply' first to create drafts.")
        return

    print(f"\nTarget font: {target_font} {target_size}pt")
    print(f"Mode: {'APPLY (will modify files)' if apply else 'PREVIEW (no changes)'}")
    print(f"Documents: {len(drafts)}")

    if apply:
        # Create backup before applying changes
        print("\n" + "-" * 40)
        print("CREATING BACKUP")
        print("-" * 40)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_subdir = config.backup_dir / f"pre_font_standardize_{timestamp}"
        backup_subdir.mkdir(parents=True, exist_ok=True)

        for draft_path in drafts:
            backup_path = backup_subdir / draft_path.name
            shutil.copy2(draft_path, backup_path)
            print(f"  Backed up: {draft_path.name}")

        print(f"\nBackups saved to: {backup_subdir}")

    print("\n" + "-" * 40)
    print("PROCESSING DOCUMENTS")
    print("-" * 40)

    standardizer = FontStandardizer(config, target_font, target_size)
    results = standardizer.standardize_drafts(preview_only=not apply)

    # Generate reports
    config.ensure_dirs()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    mode_suffix = "applied" if apply else "preview"

    # JSON
    json_path = config.output_dir / f"font_standardize_{mode_suffix}_{timestamp}.json"
    with open(json_path, 'w') as f:
        json.dump(results, f, indent=2, default=str)
    print(f"\nJSON results: {json_path}")

    # Text report
    report_lines = []
    report_lines.append("=" * 80)
    report_lines.append(f"FONT STANDARDIZATION REPORT - {results['mode'].upper()}")
    report_lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report_lines.append("=" * 80)
    report_lines.append(f"\nTarget font: {target_font} {target_size}pt")
    report_lines.append(f"Documents processed: {len(results['documents'])}")
    report_lines.append(f"Total changes: {results['total_changes']}")

    for doc_name, doc_data in results["documents"].items():
        report_lines.append(f"\n" + "-" * 40)
        report_lines.append(f"Document: {doc_name}")
        report_lines.append(f"Changes: {doc_data['change_count']}")

        if doc_data.get("changes"):
            # Group changes by old font
            changes_by_font = defaultdict(list)
            for change in doc_data["changes"]:
                before = change.get("before", {})
                font_name = before.get("font_name", "[inherited]") or "[inherited]"
                font_size = before.get("font_size_pt")
                size_str = f"{font_size}pt" if font_size else "[inherited]"
                key = f"{font_name}, {size_str}"
                changes_by_font[key].append(change)

            for old_font, changes in sorted(changes_by_font.items(), key=lambda x: -len(x[1])):
                report_lines.append(f"\n  {old_font} -> {target_font} {target_size}pt: {len(changes)} instances")
                # Show a few samples
                for change in changes[:3]:
                    sample = change.get("text_sample", "")[:50]
                    if sample:
                        report_lines.append(f"    Sample: \"{sample}...\"")

    report_lines.append("\n" + "=" * 80)
    report_lines.append("END OF REPORT")
    report_lines.append("=" * 80)

    report_path = config.output_dir / f"font_standardize_{mode_suffix}_{timestamp}.txt"
    with open(report_path, 'w') as f:
        f.write("\n".join(report_lines))
    print(f"Text report: {report_path}")

    # Summary
    print("\n" + "-" * 40)
    print("SUMMARY")
    print("-" * 40)
    print(f"\nTotal changes: {results['total_changes']}")

    if not apply and results['total_changes'] > 0:
        print("\nTo apply these changes, run:")
        print(f"  python package_manager.py standardize-fonts --apply")
    elif apply:
        print("\nFont standardization complete!")
        print("\nNext step: Run 'python package_manager.py verify' to validate the drafts.")


def main():
    parser = argparse.ArgumentParser(
        description="FedRAMP Package Manager - Manage documentation updates"
    )
    parser.add_argument(
        "command",
        choices=["status", "analyze", "preview", "apply", "verify", "export", "fonts", "standardize-fonts"],
        help="Command to run"
    )
    parser.add_argument(
        "--dir", "-d",
        default=".",
        help="Base directory for package documents"
    )
    parser.add_argument(
        "--format", "-f",
        choices=["excel", "csv"],
        default="excel",
        help="Export format (default: excel)"
    )
    parser.add_argument(
        "--apply",
        action="store_true",
        help="Apply changes (for standardize-fonts command)"
    )
    parser.add_argument(
        "--target-font",
        default="Arial",
        help="Target font name for standardization (default: Arial)"
    )
    parser.add_argument(
        "--target-size",
        type=float,
        default=11.0,
        help="Target font size in points for standardization (default: 11.0)"
    )

    args = parser.parse_args()

    config = Config(args.dir)

    if args.command == "status":
        cmd_status(config)
    elif args.command == "analyze":
        cmd_analyze(config, args.format)
    elif args.command == "preview":
        cmd_preview(config, args.format)
    elif args.command == "apply":
        cmd_apply(config, args.format)
    elif args.command == "verify":
        cmd_verify(config, args.format)
    elif args.command == "export":
        cmd_export(config, args.format)
    elif args.command == "fonts":
        cmd_fonts(config, args.format)
    elif args.command == "standardize-fonts":
        cmd_standardize_fonts(config, args.apply, args.target_font, args.target_size)


if __name__ == "__main__":
    main()
