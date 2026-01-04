#!/usr/bin/env python3
"""
FedRAMP Package Analyzer

Analyzes all Word documents in a FedRAMP package to:
1. Find all instances of known terms (technologies, teams, positions)
2. Discover potential new terms not in your dictionary
3. Generate cross-document consistency reports
4. Prepare for bulk replacements

Usage:
    python package_analyzer.py                    # Analyze all .docx files in current directory
    python package_analyzer.py --dir /path/to/docs
    python package_analyzer.py --discover         # Focus on discovering new terms
    python package_analyzer.py --report           # Generate detailed report only
"""

import re
import json
import argparse
from datetime import datetime
from pathlib import Path
from collections import defaultdict
from docx import Document


class TermsDictionary:
    """Manages the master terms dictionary."""

    def __init__(self, dict_path: str = None):
        self.dict_path = dict_path or Path(__file__).parent / "terms_dictionary.json"
        self.data = self._load()

    def _load(self) -> dict:
        """Load the terms dictionary from JSON."""
        if not Path(self.dict_path).exists():
            print(f"Warning: Terms dictionary not found at {self.dict_path}")
            return {"known_technologies": {"terms": {}},
                    "known_teams": {"terms": {}},
                    "known_positions": {"terms": {}},
                    "discovery_patterns": {},
                    "exclusions": {"terms": []}}

        with open(self.dict_path, 'r') as f:
            return json.load(f)

    def get_all_known_terms(self) -> dict:
        """Get all known terms with their categories."""
        terms = {}

        for term, info in self.data.get("known_technologies", {}).get("terms", {}).items():
            terms[term] = {"type": "technology", "category": info.get("category", "unknown")}

        for term, info in self.data.get("known_teams", {}).get("terms", {}).items():
            terms[term] = {"type": "team"}

        for term, info in self.data.get("known_positions", {}).get("terms", {}).items():
            terms[term] = {"type": "position", "acronym": info.get("acronym"),
                          "full_form": info.get("full_form")}

        return terms

    def get_discovery_patterns(self) -> dict:
        """Get regex patterns for term discovery."""
        return self.data.get("discovery_patterns", {})

    def get_exclusions(self) -> list:
        """Get list of terms to exclude from discovery."""
        return self.data.get("exclusions", {}).get("terms", [])


class DocumentAnalyzer:
    """Analyzes a single Word document."""

    def __init__(self, doc_path: str, terms_dict: TermsDictionary):
        self.doc_path = Path(doc_path)
        self.terms_dict = terms_dict
        self.doc = Document(str(doc_path))
        self.full_text = self._extract_text()

    def _extract_text(self) -> str:
        """Extract all text from the document."""
        all_text = []

        # Body paragraphs
        for para in self.doc.paragraphs:
            all_text.append(para.text)

        # Tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_text.append(para.text)

        # Headers and footers
        for section in self.doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        all_text.append(para.text)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        all_text.append(para.text)

        return "\n".join(all_text)

    def find_known_terms(self) -> dict:
        """Find all instances of known terms in the document."""
        results = {}
        known_terms = self.terms_dict.get_all_known_terms()

        for term, info in known_terms.items():
            # Use word boundary matching
            pattern = re.compile(rf'\b{re.escape(term)}\b')
            matches = pattern.findall(self.full_text)

            if matches:
                results[term] = {
                    "count": len(matches),
                    "type": info.get("type"),
                    "category": info.get("category"),
                    "contexts": self._get_contexts(term, max_contexts=3)
                }

        return results

    def _get_contexts(self, term: str, max_contexts: int = 3) -> list:
        """Get sample contexts where a term appears."""
        contexts = []
        pattern = re.compile(rf'.{{0,50}}\b{re.escape(term)}\b.{{0,50}}', re.IGNORECASE)

        for match in pattern.finditer(self.full_text):
            context = match.group().strip()
            # Clean up the context
            context = re.sub(r'\s+', ' ', context)
            contexts.append(context)
            if len(contexts) >= max_contexts:
                break

        return contexts

    def discover_potential_terms(self) -> dict:
        """Discover potential new terms not in the dictionary."""
        discovered = defaultdict(lambda: {"count": 0, "contexts": [], "source_patterns": []})
        patterns = self.terms_dict.get_discovery_patterns()
        exclusions = set(self.terms_dict.get_exclusions())
        known_terms = set(self.terms_dict.get_all_known_terms().keys())

        # Technology patterns
        for pattern_str in patterns.get("technology_indicators", []):
            try:
                pattern = re.compile(pattern_str)
                for match in pattern.finditer(self.full_text):
                    term = match.group(1).strip() if match.groups() else match.group().strip()
                    if term and term not in exclusions and term not in known_terms:
                        if len(term) > 2:  # Skip very short matches
                            discovered[term]["count"] += 1
                            discovered[term]["source_patterns"].append("technology")
                            if len(discovered[term]["contexts"]) < 2:
                                context = self.full_text[max(0, match.start()-30):match.end()+30]
                                discovered[term]["contexts"].append(context.strip())
            except re.error:
                continue

        # Team patterns
        for pattern_str in patterns.get("team_indicators", []):
            try:
                pattern = re.compile(pattern_str)
                for match in pattern.finditer(self.full_text):
                    term = match.group(1).strip() if match.groups() else match.group().strip()
                    if term and term not in exclusions and term not in known_terms:
                        discovered[term]["count"] += 1
                        discovered[term]["source_patterns"].append("team")
                        if len(discovered[term]["contexts"]) < 2:
                            context = self.full_text[max(0, match.start()-30):match.end()+30]
                            discovered[term]["contexts"].append(context.strip())
            except re.error:
                continue

        # Position patterns
        for pattern_str in patterns.get("position_indicators", []):
            try:
                pattern = re.compile(pattern_str)
                for match in pattern.finditer(self.full_text):
                    term = match.group(1).strip() if match.groups() else match.group().strip()
                    if term and term not in exclusions and term not in known_terms:
                        discovered[term]["count"] += 1
                        discovered[term]["source_patterns"].append("position")
                        if len(discovered[term]["contexts"]) < 2:
                            context = self.full_text[max(0, match.start()-30):match.end()+30]
                            discovered[term]["contexts"].append(context.strip())
            except re.error:
                continue

        # Convert to regular dict and deduplicate source patterns
        result = {}
        for term, data in discovered.items():
            result[term] = {
                "count": data["count"],
                "contexts": data["contexts"],
                "likely_type": max(set(data["source_patterns"]), key=data["source_patterns"].count)
            }

        return result


class PackageAnalyzer:
    """Analyzes all documents in a FedRAMP package."""

    def __init__(self, directory: str = ".", terms_dict_path: str = None):
        self.directory = Path(directory)
        self.terms_dict = TermsDictionary(terms_dict_path)
        self.documents = []
        self.results = {}

    def find_documents(self) -> list:
        """Find all Word documents in the directory."""
        docs = list(self.directory.glob("*.docx"))
        # Exclude temp files
        docs = [d for d in docs if not d.name.startswith("~$")]
        self.documents = sorted(docs)
        return self.documents

    def analyze_all(self, discover: bool = True) -> dict:
        """Analyze all documents in the package."""
        if not self.documents:
            self.find_documents()

        print(f"\nFound {len(self.documents)} documents to analyze:")
        for doc in self.documents:
            print(f"  - {doc.name}")

        self.results = {
            "timestamp": datetime.now().isoformat(),
            "documents_analyzed": len(self.documents),
            "known_terms": defaultdict(lambda: {"documents": [], "total_count": 0}),
            "discovered_terms": defaultdict(lambda: {"documents": [], "total_count": 0, "likely_type": None}),
            "per_document": {}
        }

        for doc_path in self.documents:
            print(f"\nAnalyzing: {doc_path.name}...")
            try:
                analyzer = DocumentAnalyzer(str(doc_path), self.terms_dict)

                # Find known terms
                known = analyzer.find_known_terms()
                for term, data in known.items():
                    self.results["known_terms"][term]["documents"].append({
                        "file": doc_path.name,
                        "count": data["count"],
                        "contexts": data["contexts"]
                    })
                    self.results["known_terms"][term]["total_count"] += data["count"]
                    self.results["known_terms"][term]["type"] = data.get("type")
                    self.results["known_terms"][term]["category"] = data.get("category")

                # Discover new terms
                if discover:
                    discovered = analyzer.discover_potential_terms()
                    for term, data in discovered.items():
                        self.results["discovered_terms"][term]["documents"].append({
                            "file": doc_path.name,
                            "count": data["count"],
                            "contexts": data["contexts"]
                        })
                        self.results["discovered_terms"][term]["total_count"] += data["count"]
                        self.results["discovered_terms"][term]["likely_type"] = data["likely_type"]

                self.results["per_document"][doc_path.name] = {
                    "known_terms_found": len(known),
                    "discovered_terms": len(discovered) if discover else 0
                }

            except Exception as e:
                print(f"  Error analyzing {doc_path.name}: {e}")
                self.results["per_document"][doc_path.name] = {"error": str(e)}

        # Convert defaultdicts to regular dicts
        self.results["known_terms"] = dict(self.results["known_terms"])
        self.results["discovered_terms"] = dict(self.results["discovered_terms"])

        return self.results

    def generate_report(self, output_path: str = None) -> str:
        """Generate a comprehensive analysis report."""
        lines = []
        lines.append("=" * 80)
        lines.append("FEDRAMP PACKAGE ANALYSIS REPORT")
        lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"Documents analyzed: {self.results.get('documents_analyzed', 0)}")
        lines.append("=" * 80)

        # Summary
        lines.append("\n" + "-" * 40)
        lines.append("SUMMARY")
        lines.append("-" * 40)

        known_count = len(self.results.get("known_terms", {}))
        discovered_count = len(self.results.get("discovered_terms", {}))

        lines.append(f"Known terms found: {known_count}")
        lines.append(f"Potential new terms discovered: {discovered_count}")

        # Known terms by category
        lines.append("\n" + "-" * 40)
        lines.append("KNOWN TERMS FOUND")
        lines.append("-" * 40)

        # Group by type
        by_type = defaultdict(list)
        for term, data in sorted(self.results.get("known_terms", {}).items(),
                                  key=lambda x: -x[1]["total_count"]):
            term_type = data.get("type", "unknown")
            by_type[term_type].append((term, data))

        for term_type, terms in sorted(by_type.items()):
            lines.append(f"\n{term_type.upper()}:")
            for term, data in terms:
                doc_count = len(data["documents"])
                lines.append(f"  [{data['total_count']:4d}x in {doc_count} doc(s)] {term}")

        # Discovered terms
        if self.results.get("discovered_terms"):
            lines.append("\n" + "-" * 40)
            lines.append("DISCOVERED TERMS (potential new additions)")
            lines.append("-" * 40)
            lines.append("\nReview these terms and add relevant ones to terms_dictionary.json:\n")

            # Sort by count
            sorted_discovered = sorted(
                self.results["discovered_terms"].items(),
                key=lambda x: -x[1]["total_count"]
            )

            # Group by likely type
            by_type = defaultdict(list)
            for term, data in sorted_discovered:
                likely_type = data.get("likely_type", "unknown")
                by_type[likely_type].append((term, data))

            for likely_type, terms in sorted(by_type.items()):
                lines.append(f"\nLikely {likely_type.upper()}:")
                for term, data in terms[:20]:  # Limit to top 20 per category
                    doc_count = len(data["documents"])
                    lines.append(f"  [{data['total_count']:4d}x in {doc_count} doc(s)] {term}")
                    if data.get("contexts"):
                        lines.append(f"       Context: \"{data['contexts'][0][:80]}...\"")

        # Per-document summary
        lines.append("\n" + "-" * 40)
        lines.append("PER-DOCUMENT SUMMARY")
        lines.append("-" * 40)

        for doc_name, doc_data in self.results.get("per_document", {}).items():
            if "error" in doc_data:
                lines.append(f"\n{doc_name}: ERROR - {doc_data['error']}")
            else:
                lines.append(f"\n{doc_name}:")
                lines.append(f"  Known terms found: {doc_data['known_terms_found']}")
                lines.append(f"  Potential new terms: {doc_data.get('discovered_terms', 0)}")

        # Cross-document consistency
        lines.append("\n" + "-" * 40)
        lines.append("CROSS-DOCUMENT TERM USAGE")
        lines.append("-" * 40)
        lines.append("\nTerms appearing in multiple documents:\n")

        multi_doc_terms = [
            (term, data) for term, data in self.results.get("known_terms", {}).items()
            if len(data["documents"]) > 1
        ]

        for term, data in sorted(multi_doc_terms, key=lambda x: -len(x[1]["documents"])):
            doc_names = [d["file"] for d in data["documents"]]
            lines.append(f"  {term}:")
            for doc in data["documents"]:
                lines.append(f"    - {doc['file']}: {doc['count']} occurrences")

        lines.append("\n" + "=" * 80)
        lines.append("END OF REPORT")
        lines.append("=" * 80)

        report = "\n".join(lines)

        if output_path:
            with open(output_path, 'w') as f:
                f.write(report)

        return report

    def generate_replacement_suggestions(self, output_path: str = None) -> str:
        """Generate a JSON file with suggested replacements based on analysis."""
        suggestions = {
            "_description": "Suggested replacements based on package analysis",
            "_generated": datetime.now().isoformat(),
            "replacements": {}
        }

        # Add all known terms found in documents
        for term, data in self.results.get("known_terms", {}).items():
            category = data.get("type", "unknown")
            suggestions["replacements"][term] = {
                "current": term,
                "replacement": None,
                "type": category,
                "occurrences": data["total_count"],
                "documents": [d["file"] for d in data["documents"]]
            }

        output = json.dumps(suggestions, indent=2)

        if output_path:
            with open(output_path, 'w') as f:
                f.write(output)

        return output


def main():
    parser = argparse.ArgumentParser(
        description="Analyze FedRAMP package documents for terms and technologies"
    )
    parser.add_argument(
        "--dir", "-d",
        default=".",
        help="Directory containing the package documents (default: current directory)"
    )
    parser.add_argument(
        "--discover",
        action="store_true",
        default=True,
        help="Discover potential new terms (default: True)"
    )
    parser.add_argument(
        "--no-discover",
        action="store_true",
        help="Skip term discovery (faster)"
    )
    parser.add_argument(
        "--report-only",
        action="store_true",
        help="Generate report without re-analyzing"
    )
    parser.add_argument(
        "--output", "-o",
        help="Output file for the report"
    )

    args = parser.parse_args()

    print("\n" + "=" * 60)
    print("FEDRAMP PACKAGE ANALYZER")
    print("=" * 60)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    analyzer = PackageAnalyzer(args.dir)
    docs = analyzer.find_documents()

    if not docs:
        print(f"\nNo .docx files found in {args.dir}")
        return

    # Analyze documents
    discover = not args.no_discover
    results = analyzer.analyze_all(discover=discover)

    # Generate report
    report_path = args.output or f"package_analysis_{timestamp}.txt"
    report = analyzer.generate_report(report_path)
    print(f"\nReport saved to: {report_path}")

    # Generate JSON results
    json_path = f"package_analysis_{timestamp}.json"
    with open(json_path, 'w') as f:
        json.dump(results, f, indent=2, default=str)
    print(f"JSON results saved to: {json_path}")

    # Generate replacement suggestions
    suggestions_path = f"suggested_replacements_{timestamp}.json"
    analyzer.generate_replacement_suggestions(suggestions_path)
    print(f"Replacement suggestions saved to: {suggestions_path}")

    # Print summary
    print("\n" + "-" * 40)
    print("QUICK SUMMARY")
    print("-" * 40)
    print(f"Documents analyzed: {len(docs)}")
    print(f"Known terms found: {len(results.get('known_terms', {}))}")
    print(f"Potential new terms: {len(results.get('discovered_terms', {}))}")

    # Show top discovered terms
    if results.get("discovered_terms"):
        print("\nTop 10 discovered terms (consider adding to dictionary):")
        sorted_discovered = sorted(
            results["discovered_terms"].items(),
            key=lambda x: -x[1]["total_count"]
        )[:10]
        for term, data in sorted_discovered:
            print(f"  [{data['total_count']:3d}x] {term} (likely {data['likely_type']})")

    print("\n" + "=" * 60)
    print("COMPLETE")
    print("=" * 60)


if __name__ == "__main__":
    main()
