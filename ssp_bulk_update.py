#!/usr/bin/env python3
"""
SSP Bulk Update Tool

This script performs bulk find-and-replace operations on a Word document (.docx)
while preserving formatting and generating a detailed audit trail.

Usage:
    1. Edit replacements.json with your old -> new mappings
    2. Run: python ssp_bulk_update.py --preview    (to review changes first)
    3. Run: python ssp_bulk_update.py              (to apply changes)
    4. Review the change log and validation report
"""

import re
import json
import argparse
from datetime import datetime
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


# =============================================================================
# CONFIGURATION - Edit these values
# =============================================================================

# Input document path
INPUT_FILE = "Appendix A_FedRAMP Mod_ Security Controls_Version 2.3.docx"

# Output file (will be created with timestamp if not specified)
OUTPUT_FILE = None  # Set to a filename or leave as None for auto-naming

# Replacements can be defined here OR loaded from replacements.json
# If replacements.json exists, it will be used instead of this dictionary
REPLACEMENTS = {
    # ===================
    # TECHNOLOGY CHANGES
    # ===================
    # Example: Uncomment and modify as needed
    # "Wiz": "CrowdStrike Falcon",
    # "BurpSuite Enterprise": "Veracode DAST",
    # "Jira": "ServiceNow",
    # "PagerDuty": "Opsgenie",

    # ===================
    # TEAM NAME CHANGES
    # ===================
    # Example: Uncomment and modify as needed
    # "FMSP Operations": "Cloud Infrastructure Team",
    # "FMSP Engineering": "Platform Engineering Team",
    # "FMSP Security Office": "Information Security Team",
    # "Technical Operations": "Site Reliability Engineering",

    # ===================
    # POSITION CHANGES
    # ===================
    # Example: Uncomment and modify as needed
    # "Chief Information Security Officer": "Chief Security Officer",
}

# Use whole-word matching to avoid partial replacements
# e.g., "AWS" won't match "LAWS"
WHOLE_WORD_MATCHING = True

# Case-sensitive matching
CASE_SENSITIVE = True


def load_replacements_from_json(json_path: str) -> dict:
    """Load replacements from a JSON configuration file.

    Special values:
    - "DELETE" or "REMOVE": Removes the term entirely (replaces with empty string)
    - "DELETE_THIS_LINE": Skips this entry (used for section headers)
    - Empty string "": Also removes the term
    """
    with open(json_path, 'r') as f:
        data = json.load(f)

    replacements = data.get("replacements", {})

    # Filter out instruction/separator lines
    filtered = {}
    for k, v in replacements.items():
        # Skip section headers and instruction lines
        if k.startswith("=====") or v == "DELETE_THIS_LINE" or k.startswith("_"):
            continue

        # Convert DELETE/REMOVE keywords to empty string
        if v in ("DELETE", "REMOVE"):
            filtered[k] = ""
        else:
            filtered[k] = v

    return filtered


# =============================================================================
# SCRIPT LOGIC - No need to edit below this line
# =============================================================================

class ChangeLog:
    """Tracks all changes made during document processing."""

    def __init__(self):
        self.changes = []
        self.summary = {}

    def add_change(self, old_text: str, new_text: str, location: str,
                   context_before: str, context_after: str):
        """Record a single change with before/after context."""
        is_deletion = new_text == ""

        self.changes.append({
            "old_text": old_text,
            "new_text": new_text,
            "location": location,
            "context_before": context_before[:300] + "..." if len(context_before) > 300 else context_before,
            "context_after": context_after[:300] + "..." if len(context_after) > 300 else context_after,
            "is_deletion": is_deletion
        })

        if is_deletion:
            key = f"{old_text} -> [DELETED]"
        else:
            key = f"{old_text} -> {new_text}"
        self.summary[key] = self.summary.get(key, 0) + 1

    def generate_report(self) -> str:
        """Generate a formatted change report."""
        lines = []
        lines.append("=" * 80)
        lines.append("SSP BULK UPDATE - CHANGE REPORT")
        lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append("=" * 80)

        lines.append("\n" + "-" * 40)
        lines.append("SUMMARY")
        lines.append("-" * 40)

        if not self.summary:
            lines.append("No changes were made.")
        else:
            total = sum(self.summary.values())
            lines.append(f"Total replacements: {total}\n")

            for change, count in sorted(self.summary.items(), key=lambda x: -x[1]):
                lines.append(f"  [{count:4d}x] {change}")

        lines.append("\n" + "-" * 40)
        lines.append("DETAILED CHANGES")
        lines.append("-" * 40)

        if not self.changes:
            lines.append("No changes were made.")
        else:
            for i, change in enumerate(self.changes, 1):
                lines.append(f"\n[{i}] {change['location']}")
                if change.get('is_deletion'):
                    lines.append(f"    DELETION: '{change['old_text']}' [REMOVED]")
                else:
                    lines.append(f"    REPLACEMENT: '{change['old_text']}' -> '{change['new_text']}'")
                lines.append(f"    BEFORE: ...{change['context_before']}...")
                lines.append(f"    AFTER:  ...{change['context_after']}...")

        lines.append("\n" + "=" * 80)
        lines.append("END OF REPORT")
        lines.append("=" * 80)

        return "\n".join(lines)

    def generate_validation_report(self) -> str:
        """Generate a validation report for reviewing context accuracy."""
        lines = []
        lines.append("=" * 80)
        lines.append("SSP BULK UPDATE - VALIDATION REPORT")
        lines.append("Review each change to ensure the context reads correctly.")
        lines.append("=" * 80)

        if not self.changes:
            lines.append("\nNo changes to validate.")
            return "\n".join(lines)

        # Group changes by replacement type
        by_replacement = {}
        for change in self.changes:
            if change.get('is_deletion'):
                key = f"{change['old_text']} -> [DELETED]"
            else:
                key = f"{change['old_text']} -> {change['new_text']}"
            if key not in by_replacement:
                by_replacement[key] = []
            by_replacement[key].append(change)

        for replacement, changes in sorted(by_replacement.items()):
            lines.append(f"\n{'=' * 60}")
            lines.append(f"REPLACEMENT: {replacement}")
            lines.append(f"Total occurrences: {len(changes)}")
            lines.append("=" * 60)

            for i, change in enumerate(changes, 1):
                lines.append(f"\n--- [{i}/{len(changes)}] {change['location']} ---")
                lines.append("BEFORE:")
                lines.append(f"  \"{change['context_before']}\"")
                lines.append("AFTER:")
                lines.append(f"  \"{change['context_after']}\"")
                lines.append("")

        lines.append("\n" + "=" * 80)
        lines.append("END OF VALIDATION REPORT")
        lines.append("=" * 80)

        return "\n".join(lines)

    def save_json(self, filepath: str):
        """Save changes as JSON for programmatic access."""
        data = {
            "generated": datetime.now().isoformat(),
            "summary": self.summary,
            "total_changes": len(self.changes),
            "changes": self.changes
        }
        with open(filepath, 'w') as f:
            json.dump(data, f, indent=2)


def create_pattern(text: str, whole_word: bool, case_sensitive: bool) -> re.Pattern:
    """Create a regex pattern for the search text."""
    escaped = re.escape(text)
    if whole_word:
        pattern = rf'\b{escaped}\b'
    else:
        pattern = escaped

    flags = 0 if case_sensitive else re.IGNORECASE
    return re.compile(pattern, flags)


# Common acronyms and their full forms for consistency checking
ACRONYM_MAP = {
    "Chief Information Security Officer": "CISO",
    "Information System Security Officer": "ISSO",
    "Information Security": "IS",
    "System Administrator": "SA",
    "System Security Plan": "SSP",
    "Plan of Action and Milestones": "POA&M",
    "Security Assessment Report": "SAR",
    "Security Assessment Plan": "SAP",
    "Access Control": "AC",
    "Audit and Accountability": "AU",
    "Configuration Management": "CM",
    "Contingency Planning": "CP",
    "Identification and Authentication": "IA",
    "Incident Response": "IR",
    "Risk Assessment": "RA",
    "System and Communications Protection": "SC",
    "System and Information Integrity": "SI",
    "Program Manager": "PM",
    "Human Resources": "HR",
    "Information Technology": "IT",
    "Multi-Factor Authentication": "MFA",
    "Single Sign-On": "SSO",
    "Virtual Private Network": "VPN",
    "Dynamic Application Security Testing": "DAST",
    "Static Application Security Testing": "SAST",
    "Security Information and Event Management": "SIEM",
}


def check_acronym_consistency(replacements: dict, document_text: str,
                               whole_word: bool, case_sensitive: bool) -> list:
    """
    Check for acronym consistency issues.

    Returns a list of warnings about acronyms that exist in the document
    but are not being replaced when their full form is being replaced.
    """
    warnings = []

    for old_text, new_text in replacements.items():
        # Check if this replacement has a known acronym
        if old_text in ACRONYM_MAP:
            acronym = ACRONYM_MAP[old_text]

            # Check if the acronym exists in the document
            acronym_pattern = create_pattern(acronym, whole_word, case_sensitive)
            if acronym_pattern.search(document_text):
                # Check if the acronym is already being replaced
                acronym_being_replaced = any(
                    k == acronym or k.startswith(acronym + " ") or k.endswith(" " + acronym)
                    for k in replacements.keys()
                )

                if not acronym_being_replaced:
                    # Count occurrences
                    count = len(acronym_pattern.findall(document_text))
                    warnings.append({
                        "type": "acronym",
                        "full_form": old_text,
                        "acronym": acronym,
                        "count": count,
                        "message": f"'{old_text}' is being replaced, but '{acronym}' ({count} occurrences) is not. Consider adding: \"{acronym}\": \"YOUR_NEW_ACRONYM\""
                    })

    return warnings


def check_plural_forms(replacements: dict, document_text: str,
                       whole_word: bool, case_sensitive: bool) -> list:
    """
    Check for plural forms that might be missed.

    Returns a list of warnings about plural forms that exist in the document
    but are not explicitly defined in replacements.
    """
    warnings = []

    for old_text, new_text in replacements.items():
        # Skip if this is already a plural or ends with common plural suffixes
        if old_text.endswith('s') or old_text.endswith('es') or old_text.endswith('ies'):
            continue

        # Generate possible plural forms
        plural_forms = []

        # Standard plurals
        if old_text.endswith('y') and len(old_text) > 1 and old_text[-2] not in 'aeiou':
            # e.g., "Policy" -> "Policies"
            plural_forms.append(old_text[:-1] + 'ies')
        elif old_text.endswith(('s', 'x', 'z', 'ch', 'sh')):
            # e.g., "Process" -> "Processes"
            plural_forms.append(old_text + 'es')
        else:
            # Standard: add 's'
            plural_forms.append(old_text + 's')

        # Check each plural form
        for plural in plural_forms:
            # Skip if plural is already being replaced
            if plural in replacements:
                continue

            plural_pattern = create_pattern(plural, whole_word, case_sensitive)
            if plural_pattern.search(document_text):
                count = len(plural_pattern.findall(document_text))
                # Suggest the plural replacement
                if new_text == "":
                    suggested_new = "DELETE"
                elif new_text.endswith('y') and len(new_text) > 1 and new_text[-2] not in 'aeiou':
                    suggested_new = new_text[:-1] + 'ies'
                elif new_text.endswith(('s', 'x', 'z', 'ch', 'sh')):
                    suggested_new = new_text + 'es'
                else:
                    suggested_new = new_text + 's'

                warnings.append({
                    "type": "plural",
                    "singular": old_text,
                    "plural": plural,
                    "count": count,
                    "suggested_replacement": suggested_new,
                    "message": f"Plural form '{plural}' ({count} occurrences) found but not in replacements. Consider adding: \"{plural}\": \"{suggested_new}\""
                })

    return warnings


def extract_document_text(doc) -> str:
    """Extract all text from a document for analysis."""
    all_text = []

    # Body paragraphs
    for para in doc.paragraphs:
        all_text.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_text.append(para.text)

    # Headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    all_text.append(para.text)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    all_text.append(para.text)

    return "\n".join(all_text)


def order_replacements(replacements: dict) -> list:
    """
    Order replacements so longer strings are processed first.

    This prevents partial replacements, e.g., "FMSP Security" replacing
    part of "FMSP Security Office" before it can be matched.

    Returns a list of (old_text, new_text) tuples sorted by length descending.
    """
    return sorted(replacements.items(), key=lambda x: len(x[0]), reverse=True)


def replace_in_paragraph(paragraph: Paragraph, replacements: dict,
                         change_log: ChangeLog, location: str,
                         whole_word: bool, case_sensitive: bool,
                         preview_only: bool = False) -> bool:
    """
    Replace text in a paragraph while preserving formatting.
    Returns True if any changes were made.

    If preview_only=True, only logs changes without modifying the document.

    Formatting preservation strategy:
    - Build a map of character positions to runs
    - For each match, identify which runs contain it
    - Replace text within runs, preserving their formatting properties
    """
    full_text = paragraph.text
    if not full_text.strip():
        return False

    changes_made = False

    # Order replacements by length (longest first) to prevent partial matches
    ordered_replacements = order_replacements(replacements)

    # Calculate what the text will look like after all replacements
    new_full_text = full_text
    for old_text, new_text in ordered_replacements:
        pattern = create_pattern(old_text, whole_word, case_sensitive)
        new_full_text = pattern.sub(new_text, new_full_text)

    # Log each occurrence with before/after context
    # Use ordered replacements for consistent logging
    for old_text, new_text in ordered_replacements:
        pattern = create_pattern(old_text, whole_word, case_sensitive)

        if pattern.search(full_text):
            for match in pattern.finditer(full_text):
                # Get context window around the match
                context_start = max(0, match.start() - 50)
                context_end = min(len(full_text), match.end() + 50)

                # Context before replacement
                context_before = full_text[context_start:context_end]

                # Calculate the same window in the new text
                # Account for length changes from previous replacements in the same text
                before_match = full_text[:match.start()]
                new_before_match = before_match
                for ot, nt in ordered_replacements:
                    p = create_pattern(ot, whole_word, case_sensitive)
                    new_before_match = p.sub(nt, new_before_match)

                new_start = len(new_before_match)
                new_match_end = new_start + len(new_text)
                new_context_start = max(0, new_start - 50)
                new_context_end = min(len(new_full_text), new_match_end + 50)

                context_after = new_full_text[new_context_start:new_context_end]

                change_log.add_change(old_text, new_text, location,
                                     context_before, context_after)

            changes_made = True

    # Only modify if not in preview mode
    if changes_made and not preview_only:
        # Formatting-preserving replacement strategy:
        # Replace text within each run individually to preserve run formatting
        if paragraph.runs:
            # Build a list of (run_index, start_pos, end_pos) for each run
            run_positions = []
            current_pos = 0
            for i, run in enumerate(paragraph.runs):
                run_len = len(run.text)
                run_positions.append((i, current_pos, current_pos + run_len))
                current_pos += run_len

            # Apply replacements to each run's text individually
            # This preserves the formatting of each run
            for i, run in enumerate(paragraph.runs):
                run_text = run.text
                for old_text, new_text in ordered_replacements:
                    pattern = create_pattern(old_text, whole_word, case_sensitive)
                    run_text = pattern.sub(new_text, run_text)
                run.text = run_text
        else:
            paragraph.text = new_full_text

    return changes_made


def replace_in_table(table: Table, replacements: dict, change_log: ChangeLog,
                     table_num: int, whole_word: bool, case_sensitive: bool,
                     preview_only: bool = False) -> bool:
    """Replace text in all cells of a table."""
    changes_made = False

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            for para_idx, paragraph in enumerate(cell.paragraphs):
                location = f"Table {table_num}, Row {row_idx + 1}, Col {col_idx + 1}"
                if replace_in_paragraph(paragraph, replacements, change_log,
                                        location, whole_word, case_sensitive,
                                        preview_only):
                    changes_made = True

    return changes_made


def process_document(input_path: str, output_path: str, replacements: dict,
                     whole_word: bool = True, case_sensitive: bool = True,
                     preview_only: bool = False) -> ChangeLog:
    """
    Process the document, performing all replacements.
    Returns the change log.

    If preview_only=True, scans the document and logs changes without modifying.
    """
    print(f"Loading document: {input_path}")
    doc = Document(input_path)
    change_log = ChangeLog()

    # Process paragraphs in the main body
    print("Processing document body...")
    for para_idx, paragraph in enumerate(doc.paragraphs):
        location = f"Body, Paragraph {para_idx + 1}"
        replace_in_paragraph(paragraph, replacements, change_log,
                           location, whole_word, case_sensitive, preview_only)

    # Process tables
    print(f"Processing {len(doc.tables)} tables...")
    for table_idx, table in enumerate(doc.tables):
        replace_in_table(table, replacements, change_log, table_idx + 1,
                        whole_word, case_sensitive, preview_only)

    # Process headers and footers
    print("Processing headers and footers...")
    for section_idx, section in enumerate(doc.sections):
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para_idx, paragraph in enumerate(header.paragraphs):
                    location = f"Header (Section {section_idx + 1}), Paragraph {para_idx + 1}"
                    replace_in_paragraph(paragraph, replacements, change_log,
                                       location, whole_word, case_sensitive, preview_only)

        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para_idx, paragraph in enumerate(footer.paragraphs):
                    location = f"Footer (Section {section_idx + 1}), Paragraph {para_idx + 1}"
                    replace_in_paragraph(paragraph, replacements, change_log,
                                       location, whole_word, case_sensitive, preview_only)

    # Only save if not in preview mode
    if not preview_only:
        print(f"Saving updated document: {output_path}")
        doc.save(output_path)
    else:
        print("Preview mode - no changes saved to document")

    return change_log


def main():
    """Main entry point."""
    # Parse command line arguments
    parser = argparse.ArgumentParser(
        description="SSP Bulk Update Tool - Find and replace text in Word documents"
    )
    parser.add_argument(
        "--preview", "-p",
        action="store_true",
        help="Preview changes without modifying the document"
    )
    args = parser.parse_args()

    preview_mode = args.preview

    print("\n" + "=" * 60)
    if preview_mode:
        print("SSP BULK UPDATE TOOL - PREVIEW MODE")
    else:
        print("SSP BULK UPDATE TOOL")
    print("=" * 60)

    # Try to load replacements from JSON file first
    json_config_path = Path(__file__).parent / "replacements.json"
    if json_config_path.exists():
        print(f"\nLoading replacements from: {json_config_path.name}")
        replacements = load_replacements_from_json(str(json_config_path))
    else:
        print("\nUsing replacements defined in script")
        replacements = REPLACEMENTS

    # Validate replacements
    if not replacements:
        print("\nERROR: No replacements defined!")
        print("Please either:")
        print("  1. Edit replacements.json, or")
        print("  2. Edit the REPLACEMENTS dictionary in this script")
        return

    # Check input file exists
    input_path = Path(INPUT_FILE)
    if not input_path.exists():
        print(f"ERROR: Input file not found: {INPUT_FILE}")
        return

    # Generate output filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    if OUTPUT_FILE:
        output_path = Path(OUTPUT_FILE)
    else:
        output_path = input_path.parent / f"{input_path.stem}_UPDATED_{timestamp}{input_path.suffix}"

    # Create backup only if not in preview mode
    if not preview_mode:
        backup_path = input_path.parent / f"{input_path.stem}_BACKUP_{timestamp}{input_path.suffix}"
        print(f"\nCreating backup: {backup_path.name}")
        import shutil
        shutil.copy2(input_path, backup_path)
    else:
        backup_path = None
        print("\nPreview mode - no backup created")

    print(f"\nReplacements to apply: {len(replacements)}")
    for old, new in replacements.items():
        if new == "":
            print(f"  '{old}' -> [DELETE]")
        else:
            print(f"  '{old}' -> '{new}'")

    print(f"\nWhole-word matching: {WHOLE_WORD_MATCHING}")
    print(f"Case-sensitive: {CASE_SENSITIVE}")

    # Run consistency checks before processing
    print("\n" + "-" * 40)
    print("Running consistency checks...")
    print("-" * 40)

    # Load document for analysis
    doc_for_analysis = Document(str(input_path))
    document_text = extract_document_text(doc_for_analysis)

    # Check for acronym consistency
    acronym_warnings = check_acronym_consistency(
        replacements, document_text, WHOLE_WORD_MATCHING, CASE_SENSITIVE
    )

    # Check for plural forms
    plural_warnings = check_plural_forms(
        replacements, document_text, WHOLE_WORD_MATCHING, CASE_SENSITIVE
    )

    all_warnings = acronym_warnings + plural_warnings

    if all_warnings:
        print(f"\n⚠️  Found {len(all_warnings)} potential consistency issues:\n")

        # Group by type
        acronym_issues = [w for w in all_warnings if w['type'] == 'acronym']
        plural_issues = [w for w in all_warnings if w['type'] == 'plural']

        if acronym_issues:
            print("ACRONYM ISSUES:")
            for w in acronym_issues:
                print(f"  • {w['message']}")
            print()

        if plural_issues:
            print("PLURAL FORM ISSUES:")
            for w in plural_issues:
                print(f"  • {w['message']}")
            print()

        # Save warnings to a file
        warnings_path = input_path.parent / f"consistency_warnings_{timestamp}.txt"
        with open(warnings_path, 'w') as f:
            f.write("SSP BULK UPDATE - CONSISTENCY WARNINGS\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 60 + "\n\n")

            if acronym_issues:
                f.write("ACRONYM ISSUES:\n")
                f.write("-" * 40 + "\n")
                for w in acronym_issues:
                    f.write(f"\nFull form: {w['full_form']}\n")
                    f.write(f"Acronym: {w['acronym']} ({w['count']} occurrences)\n")
                    f.write(f"Suggestion: Add \"{w['acronym']}\": \"YOUR_NEW_ACRONYM\" to replacements.json\n")
                f.write("\n")

            if plural_issues:
                f.write("PLURAL FORM ISSUES:\n")
                f.write("-" * 40 + "\n")
                for w in plural_issues:
                    f.write(f"\nSingular: {w['singular']}\n")
                    f.write(f"Plural: {w['plural']} ({w['count']} occurrences)\n")
                    f.write(f"Suggestion: Add \"{w['plural']}\": \"{w['suggested_replacement']}\" to replacements.json\n")

        print(f"Warnings saved to: {warnings_path.name}")
        print("\nConsider updating replacements.json to address these issues.")
        print("You can proceed, but some text may be inconsistently replaced.")
    else:
        print("✓ No consistency issues found.")

    # Process the document
    print("\n" + "-" * 40)
    change_log = process_document(
        str(input_path),
        str(output_path),
        replacements,
        whole_word=WHOLE_WORD_MATCHING,
        case_sensitive=CASE_SENSITIVE,
        preview_only=preview_mode
    )

    # Generate reports
    if preview_mode:
        # In preview mode, generate validation report
        report_txt_path = input_path.parent / f"preview_validation_{timestamp}.txt"
        report_json_path = input_path.parent / f"preview_validation_{timestamp}.json"

        print(f"\nSaving validation report (text): {report_txt_path.name}")
        with open(report_txt_path, 'w') as f:
            f.write(change_log.generate_validation_report())

        print(f"Saving validation data (JSON): {report_json_path.name}")
        change_log.save_json(str(report_json_path))

        # Print summary
        print("\n" + "=" * 60)
        print("PREVIEW SUMMARY")
        print("=" * 60)
        print(f"Total replacements that WOULD be made: {len(change_log.changes)}")
        print(f"\nOutput files:")
        print(f"  Validation report: {report_txt_path.name}")
        print(f"  Validation JSON:   {report_json_path.name}")

        if change_log.summary:
            print("\nBreakdown by replacement:")
            for change, count in sorted(change_log.summary.items(), key=lambda x: -x[1]):
                print(f"  [{count:4d}x] {change}")

        print("\n" + "-" * 40)
        print("Review the validation report to verify context is correct.")
        print("When ready, run without --preview to apply changes:")
        print(f"  python {Path(__file__).name}")
        print("-" * 40)

    else:
        # Normal mode - generate change log
        report_txt_path = input_path.parent / f"change_log_{timestamp}.txt"
        report_json_path = input_path.parent / f"change_log_{timestamp}.json"

        print(f"\nSaving change log (text): {report_txt_path.name}")
        with open(report_txt_path, 'w') as f:
            f.write(change_log.generate_report())

        print(f"Saving change log (JSON): {report_json_path.name}")
        change_log.save_json(str(report_json_path))

        # Print summary
        print("\n" + "=" * 60)
        print("SUMMARY")
        print("=" * 60)
        print(f"Total replacements made: {len(change_log.changes)}")
        print(f"\nOutput files:")
        print(f"  Updated document: {output_path.name}")
        print(f"  Backup original:  {backup_path.name}")
        print(f"  Change log (TXT): {report_txt_path.name}")
        print(f"  Change log (JSON): {report_json_path.name}")

        if change_log.summary:
            print("\nBreakdown by replacement:")
            for change, count in sorted(change_log.summary.items(), key=lambda x: -x[1]):
                print(f"  [{count:4d}x] {change}")

    print("\n" + "=" * 60)
    print("COMPLETE")
    print("=" * 60)


if __name__ == "__main__":
    main()
